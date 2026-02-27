"""
GUI Channel - PyQt6 Desktop Interface
Modern Chat UI - Clean and Minimal Design
"""
import asyncio
import os
import sys
import time
import re
from datetime import datetime
from typing import List, Optional, TYPE_CHECKING
from pathlib import Path
import uuid
import threading

from loguru import logger

from .base import BaseChannel, IncomingMessage, OutgoingMessage, MessageHandler, MessageType
from ..utils.progress import progress_manager


def create_voice_button(chat_window):
    """创建语音按钮 - 按住说话"""
    from PyQt6.QtWidgets import QPushButton
    from PyQt6.QtCore import Qt
    
    class VoiceButton(QPushButton):
        def __init__(self, parent_chat_window):
            super().__init__("🎤")
            self._chat_window = parent_chat_window
            self._recording = False
            self.setToolTip("语音输入（按住说话）")
            self.setStyleSheet("""
                QPushButton {
                    background-color: #f7f7f8;
                    color: #202123;
                    border: 1px solid #d9d9e3;
                    border-radius: 12px;
                    font-size: 20px;
                }
                QPushButton:hover {
                    background-color: #e8f5e9;
                    border: 1px solid #10a37f;
                }
            """)
        
        def mousePressEvent(self, event):
            if event.button() == Qt.MouseButton.LeftButton:
                self._start_recording()
            super().mousePressEvent(event)
        
        def mouseReleaseEvent(self, event):
            if event.button() == Qt.MouseButton.LeftButton:
                self._stop_recording()
            super().mouseReleaseEvent(event)
        
        def _start_recording(self):
            if self._recording:
                return
            self._recording = True
            self.setText("⏹️")
            self.setStyleSheet("""
                QPushButton {
                    background-color: #ef4444;
                    color: white;
                    border: 2px solid #dc2626;
                    border-radius: 12px;
                    font-size: 20px;
                }
            """)
            self._chat_window._start_voice_recording()
        
        def _stop_recording(self):
            if not self._recording:
                return
            self._recording = False
            self.setText("🎤")
            self.setStyleSheet("""
                QPushButton {
                    background-color: #f7f7f8;
                    color: #202123;
                    border: 1px solid #d9d9e3;
                    border-radius: 12px;
                    font-size: 20px;
                }
                QPushButton:hover {
                    background-color: #e8f5e9;
                    border: 1px solid #10a37f;
                }
            """)
            self._chat_window._stop_voice_recording()
    
    return VoiceButton(chat_window)

if TYPE_CHECKING:
    from PyQt6.QtWidgets import QWidget


class GUIChannel(BaseChannel):
    name = "gui"

    def __init__(self):
        self._running = False
        self._message_handlers: List[MessageHandler] = []

    async def start(self) -> None:
        pass

    async def stop(self) -> None:
        self._running = False

    async def send(self, message: OutgoingMessage) -> bool:
        return True

    def on_message(self, handler: MessageHandler) -> None:
        self._message_handlers.append(handler)

    async def is_running(self) -> bool:
        return self._running


class ChatWindow:
    """Modern chat window - Clean and minimal design"""

    class HistoryLineEdit:
        """支持历史记录的输入框 - 将在__init__中继承QLineEdit"""
        pass

    def __init__(self, agent):
        from PyQt6.QtWidgets import (
            QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
            QLineEdit, QPushButton, QLabel, QListWidget, 
            QListWidgetItem, QScrollArea, QFrame, QSizePolicy,
            QSpacerItem, QSizePolicy as QSP, QComboBox, QSlider
        )
        from PyQt6.QtCore import Qt, QSize, QTimer, QThread, pyqtSignal, QObject
        from PyQt6.QtGui import QFont
        import markdown

        self.agent = agent
        self.markdown = markdown
        
        self.Qt = Qt
        self.QFont = QFont
        self.QTimer = QTimer
        self.QSize = QSize
        self.QThread = QThread
        self.pyqtSignal = pyqtSignal
        self.QSlider = QSlider
        self.QObject = QObject
        
        self.QWidget = QWidget
        self.QHBoxLayout = QHBoxLayout
        self.QVBoxLayout = QVBoxLayout
        self.QLabel = QLabel
        self.QPushButton = QPushButton
        self.QLineEdit = QLineEdit
        self.QListWidget = QListWidget
        self.QListWidgetItem = QListWidgetItem
        self.QComboBox = QComboBox
        self.QListWidget = QListWidget
        
        self.AGENTS_LIST = self._load_agents_list()
        
        class AgentPopupList(QListWidget):
            """智能体弹出列表"""
            def __init__(self, parent=None):
                super().__init__(parent)
                self.setWindowFlags(
                    Qt.WindowType.Tool | 
                    Qt.WindowType.FramelessWindowHint |
                    Qt.WindowType.WindowStaysOnTopHint
                )
                self.setAttribute(Qt.WidgetAttribute.WA_ShowWithoutActivating)
                self.setFocusPolicy(Qt.FocusPolicy.NoFocus)
                self.setStyleSheet("""
                    QListWidget {
                        background-color: #ffffff;
                        border: 1px solid #d9d9e3;
                        border-radius: 8px;
                        padding: 4px;
                    }
                    QListWidget::item {
                        padding: 8px 12px;
                        border-radius: 6px;
                    }
                    QListWidget::item:hover {
                        background-color: #f0f0f0;
                    }
                    QListWidget::item:selected {
                        background-color: #10a37f;
                        color: white;
                    }
                """)
                self.itemClicked.connect(self._on_item_clicked)
                self.selected_agent = None
                self._callback = None
            
            def _on_item_clicked(self, item):
                self.selected_agent = item.data(Qt.ItemDataRole.UserRole)
                if self._callback:
                    self._callback(self.selected_agent)
                self.hide()
            
            def set_callback(self, callback):
                self._callback = callback
            
            def show_agents(self, agents, filter_text=""):
                self.clear()
                filtered = []
                filter_lower = filter_text.lower()
                
                for agent_data in agents:
                    if len(agent_data) >= 5:
                        name, alias, desc, english_name, keywords = agent_data[:5]
                    else:
                        name, alias, desc = agent_data[:3]
                        english_name = ""
                        keywords = []
                    
                    if not filter_text:
                        filtered.append((name, alias, desc))
                    else:
                        matched = (
                            filter_lower in name.lower() or
                            filter_lower in alias.lower() or
                            filter_lower in english_name.lower() or
                            any(filter_lower in kw.lower() for kw in keywords)
                        )
                        if matched:
                            filtered.append((name, alias, desc))
                
                for name, alias, desc in filtered:
                    item = QListWidgetItem(f"@{name} - {desc}")
                    item.setData(Qt.ItemDataRole.UserRole, name)
                    self.addItem(item)
                
                if filtered:
                    self.setCurrentRow(0)
                    self.setFixedHeight(min(len(filtered) * 40 + 8, 200))
                    return True
                return False
        
        # 定义历史记录输入框类
        class HistoryLineEdit(QLineEdit):
            """支持历史记录和@智能体选择的输入框"""
            
            def __init__(self, parent=None, agents_list=None):
                super().__init__(parent)
                self.history = []
                self.history_index = 0
                self.current_text = ""
                self.max_history = 100
                self.agents_list = agents_list or []
                self.agent_popup = AgentPopupList(self)
                self.agent_popup.set_callback(self._insert_agent)
                self._at_start_pos = -1
                self._popup_shown = False
            
            def _insert_agent(self, agent_name):
                cursor_pos = self.cursorPosition()
                text = self.text()
                at_pos = text.rfind('@', 0, cursor_pos)
                if at_pos >= 0:
                    new_text = text[:at_pos] + f"@{agent_name} " + text[cursor_pos:]
                    self.setText(new_text)
                    self.setCursorPosition(at_pos + len(agent_name) + 2)
                self._at_start_pos = -1
                self._popup_shown = False
                self.agent_popup.hide()
            
            def on_text_changed(self):
                text = self.text()
                cursor_pos = self.cursorPosition()
                at_pos = text.rfind('@', 0, cursor_pos)
                
                if at_pos >= 0:
                    space_after_at = text.find(' ', at_pos, cursor_pos)
                    if space_after_at == -1:
                        filter_text = text[at_pos + 1:cursor_pos]
                        if len(filter_text) <= 10 and not filter_text.endswith('智能体'):
                            if self.agent_popup.show_agents(self.agents_list, filter_text):
                                self._show_popup(at_pos)
                                return
                self.agent_popup.hide()
                self._at_start_pos = -1
                self._popup_shown = False
            
            def _show_popup(self, at_pos):
                input_rect = self.rect()
                global_pos = self.mapToGlobal(input_rect.bottomLeft())
                popup_width = max(300, self.width())
                
                self.agent_popup.setGeometry(
                    global_pos.x(),
                    global_pos.y() + 5,
                    popup_width,
                    self.agent_popup.height()
                )
                self.agent_popup.show()
                self._popup_shown = True
                self.setFocus()
            
            def keyPressEvent(self, event):
                if self._popup_shown and self.agent_popup.count() > 0:
                    if event.key() == Qt.Key.Key_Down:
                        current = self.agent_popup.currentRow()
                        if current < self.agent_popup.count() - 1:
                            self.agent_popup.setCurrentRow(current + 1)
                        event.accept()
                        return
                    elif event.key() == Qt.Key.Key_Up:
                        current = self.agent_popup.currentRow()
                        if current > 0:
                            self.agent_popup.setCurrentRow(current - 1)
                        event.accept()
                        return
                    elif event.key() == Qt.Key.Key_Enter or event.key() == Qt.Key.Key_Return:
                        current_item = self.agent_popup.currentItem()
                        if current_item:
                            agent_name = current_item.data(Qt.ItemDataRole.UserRole)
                            self._insert_agent(agent_name)
                        event.accept()
                        return
                    elif event.key() == Qt.Key.Key_Escape:
                        self.agent_popup.hide()
                        self._popup_shown = False
                        event.accept()
                        return
                
                if event.key() == Qt.Key.Key_Up:
                    if self.history:
                        if self.history_index == len(self.history):
                            self.current_text = self.text()
                        
                        if self.history_index > 0:
                            self.history_index -= 1
                        
                        self.setText(self.history[self.history_index])
                        self.selectAll()
                    event.accept()
                
                elif event.key() == Qt.Key.Key_Down:
                    if self.history:
                        self.history_index += 1
                        
                        if self.history_index >= len(self.history):
                            self.setText(self.current_text)
                            self.history_index = len(self.history)
                        else:
                            self.setText(self.history[self.history_index])
                            self.selectAll()
                    event.accept()
                
                elif event.key() == Qt.Key.Key_Escape:
                    if self._popup_shown:
                        self.agent_popup.hide()
                        self._popup_shown = False
                    else:
                        self.clear()
                        self.history_index = len(self.history)
                    event.accept()
                
                else:
                    super().keyPressEvent(event)
            
            def add_to_history(self, text):
                """添加文本到历史记录"""
                if not text or not text.strip():
                    return
                
                text = text.strip()
                
                # 如果与最后一条相同，不添加
                if self.history and self.history[-1] == text:
                    return
                
                self.history.append(text)
                
                # 限制历史记录数量
                if len(self.history) > self.max_history:
                    self.history.pop(0)
                
                self.history_index = len(self.history)
        
        self.HistoryLineEdit = HistoryLineEdit

        # 创建信号助手用于跨线程通信
        class SignalHelper(QObject):
            """Helper class for cross-thread communication using Qt signals"""
            show_notification_signal = pyqtSignal(str)
            play_video_signal = pyqtSignal(str, str, str)  # video_path, directory, video_name
            show_music_minimized_signal = pyqtSignal()  # 显示音乐播放器最小化条目
            show_music_player_signal = pyqtSignal()  # 显示音乐播放器窗口
            show_homeassistant_dashboard_signal = pyqtSignal()  # 显示智能家居控制面板
            append_message_signal = pyqtSignal(str, str)  # role, content
            reset_play_button_signal = pyqtSignal()  # 重置播放按钮
            auto_speak_signal = pyqtSignal(str)  # 自动播放语音
            update_token_signal = pyqtSignal(int)  # 更新Token统计
            
            def emit_notification(self, message: str):
                """Emit signal to show notification in main thread"""
                self.show_notification_signal.emit(message)
            
            def emit_play_video(self, video_path: str, directory: str, video_name: str):
                """Emit signal to play video in main thread"""
                self.play_video_signal.emit(video_path, directory, video_name)
            
            def emit_show_music_minimized(self):
                """Emit signal to show music player minimized item"""
                self.show_music_minimized_signal.emit()
            
            def emit_show_music_player(self):
                """Emit signal to show music player window"""
                self.show_music_player_signal.emit()
            
            def emit_show_homeassistant_dashboard(self):
                """Emit signal to show Home Assistant dashboard"""
                self.show_homeassistant_dashboard_signal.emit()
            
            def emit_append_message(self, role: str, content: str):
                """Emit signal to append message in main thread"""
                self.append_message_signal.emit(role, content)
            
            def emit_reset_play_button(self):
                """Emit signal to reset play button in main thread"""
                self.reset_play_button_signal.emit()
            
            def emit_auto_speak(self, text: str):
                """Emit signal for auto speak in main thread"""
                self.auto_speak_signal.emit(text)
            
            def emit_update_token(self, tokens: int):
                """Emit signal to update token count in main thread"""
                self.update_token_signal.emit(tokens)
        
        self.signal_helper = SignalHelper()
        self.signal_helper.show_notification_signal.connect(self._on_notification_signal)
        self.signal_helper.play_video_signal.connect(self._on_play_video_signal)
        self.signal_helper.show_music_minimized_signal.connect(self._show_music_minimized)
        self.signal_helper.show_music_player_signal.connect(self._show_music_player)
        self.signal_helper.show_homeassistant_dashboard_signal.connect(self._show_homeassistant_dashboard)
        self.signal_helper.append_message_signal.connect(self._on_append_message_signal)
        self.signal_helper.reset_play_button_signal.connect(self._reset_play_button)
        self.signal_helper.auto_speak_signal.connect(self._on_auto_speak)
        self.signal_helper.update_token_signal.connect(self._update_token_count)
        
        from ..tts import get_tts_manager
        get_tts_manager().set_on_stop_callback(self.signal_helper.emit_reset_play_button)

        self._init_multi_agent_system()
        
        # 注意：Master Agent 的通知回调在 _init_multi_agent_system 中设置

        self.window = QMainWindow()
        self.window.setWindowTitle("智能助理系统")
        self.window.setMinimumSize(900, 650)
        
        # 让主窗口可以访问 ChatWindow
        self.window.chat_window = self
        
        self._apply_global_styles()
        
        central = self.QWidget()
        self.window.setCentralWidget(central)
        main_layout = self.QHBoxLayout(central)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        sidebar = self._create_sidebar()
        main_layout.addWidget(sidebar)

        chat_container = self._create_chat_area()
        main_layout.addWidget(chat_container, 1)

        self._init_conversation_manager()
        self.QTimer.singleShot(100, self._load_initial_messages)
        
        self.window.closeEvent = self._on_window_close
        self.window.resizeEvent = self._on_window_resize
        self.window.changeEvent = self._on_window_change
    
    def _on_window_change(self, event):
        """Handle window state change event"""
        from PyQt6.QtCore import Qt
        from PyQt6.QtWidgets import QMainWindow
        
        if event.type() == event.Type.WindowStateChange:
            if self.window.windowState() & Qt.WindowState.WindowMinimized:
                if hasattr(self, 'input_field') and hasattr(self.input_field, 'agent_popup'):
                    self.input_field.agent_popup.hide()
                    self.input_field._popup_shown = False
            elif self.window.windowState() & Qt.WindowState.WindowActive:
                pass
        elif event.type() == event.Type.ActivationChange:
            if not self.window.isActiveWindow():
                if hasattr(self, 'input_field') and hasattr(self.input_field, 'agent_popup'):
                    self.input_field.agent_popup.hide()
                    self.input_field._popup_shown = False
        
        QMainWindow.changeEvent(self.window, event)
    
    def _on_window_resize(self, event):
        """Handle window resize event - update message item sizes"""
        from PyQt6.QtWidgets import QMainWindow
        QMainWindow.resizeEvent(self.window, event)
        self._update_message_sizes()
    
    def _update_message_sizes(self):
        """Update all message item sizes after window resize"""
        try:
            list_width = self.messages_list.width()
            max_bubble_width = max(300, list_width - 150)
            
            for i in range(self.messages_list.count()):
                item = self.messages_list.item(i)
                widget = self.messages_list.itemWidget(item)
                if widget:
                    for bubble in widget.findChildren(self.QLabel):
                        current_max = bubble.maximumWidth()
                        if current_max == 550 or current_max == 500 or current_max > max_bubble_width:
                            bubble.setMaximumWidth(max_bubble_width)
                    
                    widget.adjustSize()
                    widget.updateGeometry()
                    
                    height = widget.sizeHint().height()
                    if height < 50:
                        height = 50
                    
                    item.setSizeHint(self.QSize(list_width, height + 16))
        except Exception as e:
            logger.debug(f"Error updating message sizes: {e}")
    
    def _load_agents_list(self):
        """动态加载智能体列表"""
        try:
            from ..agents.agent_scanner import get_agent_scanner
            scanner = get_agent_scanner()
            agents = scanner.scan_agents()
            
            agents_list = []
            for name, meta in sorted(agents.items(), key=lambda x: x[1].priority):
                if not meta.hidden:
                    english_name = name.replace("_agent", "").replace("_", " ")
                    keywords = meta.keywords if hasattr(meta, 'keywords') and meta.keywords else []
                    agents_list.append((
                        f"{meta.display_name}智能体",
                        meta.display_name,
                        meta.description or f"{meta.display_name}相关功能",
                        english_name,
                        keywords
                    ))
            return agents_list
        except Exception as e:
            logger.warning(f"动态加载智能体列表失败: {e}")
            return [
                ("通讯录智能体", "通讯录", "管理联系人信息", "contact", []),
                ("音乐智能体", "音乐", "播放音乐歌曲", "music", []),
                ("视频智能体", "视频", "播放视频内容", "video", []),
                ("邮件智能体", "邮件", "发送和管理邮件", "email", []),
                ("天气智能体", "天气", "查询天气预报", "weather", []),
                ("文件智能体", "文件", "文件操作管理", "file", []),
                ("爬虫智能体", "爬虫", "网页数据抓取", "crawler", []),
                ("开发智能体", "开发", "代码开发辅助", "developer", []),
                ("系统智能体", "系统", "系统控制操作", "os", []),
                ("新闻智能体", "新闻", "获取新闻资讯", "news", []),
                ("日历智能体", "日历", "日程管理提醒", "calendar", []),
                ("PDF智能体", "PDF", "PDF文档处理", "pdf", []),
                ("下载智能体", "下载", "文件下载管理", "download", []),
                ("应用智能体", "应用", "应用程序管理", "app", []),
                ("Web服务智能体", "Web服务", "Web服务管理", "web_server", []),
                ("音频解密智能体", "音频解密", "NCM音频解密", "audio_decrypt", []),
                ("图片转换智能体", "图片转换", "图片格式转换", "image_converter", []),
            ]
    
    def refresh_agents_list(self):
        """刷新智能体列表"""
        self.AGENTS_LIST = self._load_agents_list()
        if hasattr(self, 'input_field') and self.input_field:
            self.input_field.agents_list = self.AGENTS_LIST
        logger.info(f"🔄 GUI 智能体列表已刷新: {len(self.AGENTS_LIST)} 个")
    
    def _reset_token_count(self):
        """重置Token统计"""
        self._total_tokens = 0
        self.token_label.setText("📊 Token: 0")
        try:
            from ..utils.token_counter import token_counter
            token_counter.reset()
        except Exception:
            pass
        logger.info("📊 Token 统计已重置")
    
    def _update_token_count(self, tokens: int):
        """更新Token统计"""
        self._total_tokens += tokens
        self.token_label.setText(f"📊 Token: {self._total_tokens}")
    
    def _on_window_close(self, event):
        """Handle window close event"""
        logger.info("GUI window closing")
        event.accept()

    def _apply_global_styles(self):
        """Apply global stylesheet"""
        self.window.setStyleSheet("""
            QMainWindow {
                background-color: #ffffff;
            }
            QWidget {
                font-family: "Microsoft YaHei", "Segoe UI", sans-serif;
            }
            QPushButton {
                border-radius: 6px;
                padding: 6px 12px;
                font-size: 13px;
            }
            QPushButton:hover {
                opacity: 0.9;
            }
            QLineEdit {
                border-radius: 8px;
                padding: 10px 14px;
                font-size: 14px;
            }
            QListWidget {
                border: none;
                outline: none;
            }
            QScrollArea {
                border: none;
            }
            QScrollBar:vertical {
                border: none;
                background: transparent;
                width: 6px;
                margin: 0;
            }
            QScrollBar::handle:vertical {
                background: #d0d0d0;
                border-radius: 3px;
                min-height: 30px;
            }
            QScrollBar::handle:vertical:hover {
                background: #b0b0b0;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0;
            }
        """)

    def _create_sidebar(self) -> "QWidget":
        """Create left sidebar"""
        sidebar = self.QWidget()
        sidebar.setFixedWidth(260)
        sidebar.setObjectName("sidebar")
        sidebar.setStyleSheet("""
            #sidebar {
                background-color: #f7f7f8;
                border-right: 1px solid #e5e5e5;
            }
        """)
        layout = self.QVBoxLayout(sidebar)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(8)

        clear_chat_btn = self.QPushButton("🗑️ 清空对话")
        clear_chat_btn.setStyleSheet("""
            QPushButton {
                background-color: #10a37f;
                color: white;
                border: none;
                border-radius: 6px;
                padding: 10px;
                font-size: 13px;
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: #0d8a6a;
            }
        """)
        clear_chat_btn.clicked.connect(self._clear_conversation)
        layout.addWidget(clear_chat_btn)

        # 音乐播放器最小化条目（初始隐藏）
        self.music_minimized_item = self._create_music_minimized_item()
        layout.addWidget(self.music_minimized_item)
        self.music_minimized_item.hide()

        layout.addStretch()

        settings_btn = self.QPushButton("⚙️ 设置")
        settings_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                color: #6e6e80;
                border: 1px solid #d9d9e3;
                text-align: left;
                padding: 8px 12px;
                border-radius: 6px;
            }
            QPushButton:hover {
                background-color: #ececf1;
            }
        """)
        settings_btn.clicked.connect(self._show_settings_dialog)
        layout.addWidget(settings_btn)

        return sidebar

    def _create_sidebar_item(self, icon: str, name: str, desc: str) -> "QWidget":
        """Create a sidebar feature item"""
        item = self.QWidget()
        item.setStyleSheet("""
            QWidget {
                background-color: transparent;
                border-radius: 8px;
            }
            QWidget:hover {
                background-color: #ececf1;
            }
        """)
        item.setCursor(self.Qt.CursorShape.PointingHandCursor)
        
        layout = self.QHBoxLayout(item)
        layout.setContentsMargins(10, 8, 10, 8)
        layout.setSpacing(10)
        
        icon_label = self.QLabel(icon)
        icon_label.setStyleSheet("font-size: 18px;")
        icon_label.setFixedWidth(24)
        layout.addWidget(icon_label)
        
        text_layout = self.QVBoxLayout()
        text_layout.setSpacing(2)
        
        name_label = self.QLabel(name)
        name_label.setStyleSheet("color: #202123; font-size: 13px; font-weight: 500; background: transparent;")
        text_layout.addWidget(name_label)
        
        desc_label = self.QLabel(desc)
        desc_label.setStyleSheet("color: #6e6e80; font-size: 11px; background: transparent;")
        text_layout.addWidget(desc_label)
        
        layout.addLayout(text_layout, 1)
        
        return item

    def _create_music_minimized_item(self) -> "QWidget":
        """创建音乐播放器最小化条目"""
        item = self.QWidget()
        item.setStyleSheet("""
            QWidget#musicItem {
                background-color: #f0f0f0;
                border-radius: 8px;
                border: 1px solid #e0e0e0;
            }
            QWidget#musicItem:hover {
                background-color: #e8e8e8;
            }
        """)
        item.setObjectName("musicItem")
        item.setCursor(self.Qt.CursorShape.PointingHandCursor)
        
        layout = self.QHBoxLayout(item)
        layout.setContentsMargins(10, 8, 10, 8)
        layout.setSpacing(8)
        
        icon_label = self.QLabel("🎵")
        icon_label.setStyleSheet("font-size: 20px; background: transparent;")
        layout.addWidget(icon_label)
        
        text_layout = self.QVBoxLayout()
        text_layout.setSpacing(2)
        
        self.music_minimized_title = self.QLabel("音乐播放器")
        self.music_minimized_title.setStyleSheet("color: #202123; font-size: 13px; font-weight: 500; background: transparent;")
        text_layout.addWidget(self.music_minimized_title)
        
        self.music_minimized_status = self.QLabel("正在播放...")
        self.music_minimized_status.setStyleSheet("color: #6e6e80; font-size: 11px; background: transparent;")
        text_layout.addWidget(self.music_minimized_status)
        
        layout.addLayout(text_layout, 1)
        
        # 关闭按钮
        close_btn = self.QPushButton("关闭")
        close_btn.setFixedHeight(24)
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: #eee;
                color: #333;
                border: 1px solid #ccc;
                border-radius: 4px;
                font-size: 12px;
                padding: 2px 8px;
            }
            QPushButton:hover {
                background-color: #ff4757;
                color: white;
                border-color: #ff4757;
            }
        """)
        close_btn.clicked.connect(self._stop_music_and_hide)
        layout.addWidget(close_btn)
        
        # 点击条目打开播放器窗口
        item.mousePressEvent = lambda e: self._show_music_player() if e.button() == self.Qt.MouseButton.LeftButton else None
        
        return item

    def _create_chat_area(self) -> "QWidget":
        """Create main chat area"""
        container = self.QWidget()
        container.setStyleSheet("background-color: #ffffff;")
        layout = self.QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        header = self._create_chat_header()
        layout.addWidget(header)

        from PyQt6.QtCore import QPropertyAnimation, QEasingCurve, Qt, QTimer
        from PyQt6.QtWidgets import QListWidget, QLabel

        class BubbleLabel(QLabel):
            """支持拖拽自动滚动的消息气泡标签"""
            
            def __init__(self, text, parent=None, list_widget=None):
                super().__init__(text, parent)
                self._list_widget = list_widget
            
            def setListWidget(self, list_widget):
                self._list_widget = list_widget
            
            def mouseMoveEvent(self, event):
                if self._list_widget and event.buttons() & Qt.MouseButton.LeftButton:
                    pos_in_viewport = self._list_widget.mapFromGlobal(self.mapToGlobal(event.pos()))
                    self._list_widget._check_auto_scroll(pos_in_viewport)
                super().mouseMoveEvent(event)
            
            def mouseReleaseEvent(self, event):
                if self._list_widget:
                    self._list_widget._stop_auto_scroll()
                super().mouseReleaseEvent(event)

        class SmoothScrollListWidget(QListWidget):
            """支持平滑滚动和自动滚动的列表控件"""
            
            def __init__(self, parent=None):
                super().__init__(parent)
                self._scroll_animation = None
                self._setup_smooth_scroll()
                self.setAutoFillBackground(True)
                self.setAttribute(Qt.WidgetAttribute.WA_StaticContents, True)
                palette = self.palette()
                palette.setColor(palette.ColorRole.Base, Qt.GlobalColor.white)
                palette.setColor(palette.ColorRole.Window, Qt.GlobalColor.white)
                self.setPalette(palette)
                
                self._auto_scroll_timer = QTimer(self)
                self._auto_scroll_timer.timeout.connect(self._do_auto_scroll)
                self._auto_scroll_direction = 0
                self._auto_scroll_speed = 0
            
            def _setup_smooth_scroll(self):
                """设置平滑滚动"""
                scroll_bar = self.verticalScrollBar()
                self._scroll_animation = QPropertyAnimation(scroll_bar, b"value")
                self._scroll_animation.setDuration(200)
                self._scroll_animation.setEasingCurve(QEasingCurve.Type.OutQuad)
            
            def wheelEvent(self, event):
                """重写滚轮事件实现平滑滚动"""
                delta = event.angleDelta().y()
                scroll_bar = self.verticalScrollBar()
                
                step = int(delta * 0.5)
                
                self._scroll_animation.stop()
                self._scroll_animation.setStartValue(scroll_bar.value())
                self._scroll_animation.setEndValue(
                    max(scroll_bar.minimum(), 
                        min(scroll_bar.maximum(), scroll_bar.value() - step))
                )
                self._scroll_animation.start()
                
                event.accept()
            
            def smoothScrollToBottom(self):
                """平滑滚动到底部"""
                scroll_bar = self.verticalScrollBar()
                max_value = scroll_bar.maximum()
                
                if max_value > 0:
                    self._scroll_animation.stop()
                    self._scroll_animation.setStartValue(scroll_bar.value())
                    self._scroll_animation.setEndValue(max_value)
                    self._scroll_animation.start()
            
            def viewportEvent(self, event):
                """处理视口事件，实现拖拽自动滚动"""
                if event.type() == event.Type.MouseMove:
                    if event.buttons() & Qt.MouseButton.LeftButton:
                        self._check_auto_scroll(event.pos())
                elif event.type() == event.Type.MouseButtonRelease:
                    self._stop_auto_scroll()
                return super().viewportEvent(event)
            
            def _check_auto_scroll(self, pos):
                """检查是否需要自动滚动"""
                viewport_height = self.viewport().height()
                margin = 50
                
                if pos.y() < margin:
                    self._auto_scroll_direction = -1
                    self._auto_scroll_speed = max(1, (margin - pos.y()) // 10)
                    if not self._auto_scroll_timer.isActive():
                        self._auto_scroll_timer.start(30)
                elif pos.y() > viewport_height - margin:
                    self._auto_scroll_direction = 1
                    self._auto_scroll_speed = max(1, (pos.y() - (viewport_height - margin)) // 10)
                    if not self._auto_scroll_timer.isActive():
                        self._auto_scroll_timer.start(30)
                else:
                    self._stop_auto_scroll()
            
            def _do_auto_scroll(self):
                """执行自动滚动"""
                scroll_bar = self.verticalScrollBar()
                new_value = scroll_bar.value() + self._auto_scroll_direction * self._auto_scroll_speed
                new_value = max(scroll_bar.minimum(), min(scroll_bar.maximum(), new_value))
                scroll_bar.setValue(new_value)
            
            def _stop_auto_scroll(self):
                """停止自动滚动"""
                self._auto_scroll_timer.stop()
                self._auto_scroll_direction = 0
                self._auto_scroll_speed = 0

        self.BubbleLabel = BubbleLabel
        self.messages_list = SmoothScrollListWidget()
        self.messages_list.setStyleSheet("""
            QListWidget {
                background-color: #ffffff;
                border: none;
            }
            QListWidget::item {
                background: transparent;
                padding: 0;
                margin: 0;
            }
            QListWidget::item:selected {
                background: transparent;
            }
            QScrollBar:vertical {
                background: transparent;
                width: 8px;
                margin: 0;
            }
            QScrollBar::handle:vertical {
                background: rgba(0, 0, 0, 0.2);
                border-radius: 4px;
                min-height: 30px;
            }
            QScrollBar::handle:vertical:hover {
                background: rgba(0, 0, 0, 0.35);
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0px;
            }
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {
                background: transparent;
            }
        """)
        self.messages_list.viewport().setAutoFillBackground(True)
        viewport_palette = self.messages_list.viewport().palette()
        viewport_palette.setColor(viewport_palette.ColorRole.Base, Qt.GlobalColor.white)
        viewport_palette.setColor(viewport_palette.ColorRole.Window, Qt.GlobalColor.white)
        self.messages_list.viewport().setPalette(viewport_palette)
        self.messages_list.setVerticalScrollMode(QListWidget.ScrollMode.ScrollPerPixel)
        self.messages_list.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.messages_list.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.messages_list.setSpacing(0)
        self.messages_list.setUniformItemSizes(False)
        
        layout.addWidget(self.messages_list, 1)

        input_area = self._create_input_area()
        layout.addWidget(input_area)

        return container

    def smooth_scroll_to_bottom(self):
        """平滑滚动到底部"""
        self.messages_list.smoothScrollToBottom()

    def _create_chat_header(self) -> "QWidget":
        """Create chat header"""
        header = self.QWidget()
        header.setFixedHeight(52)
        header.setStyleSheet("""
            QWidget {
                background-color: #ffffff;
                border-bottom: 1px solid #e5e5e5;
            }
        """)
        layout = self.QHBoxLayout(header)
        layout.setContentsMargins(20, 0, 20, 0)

        self.chat_title = self.QLabel("智能助手")
        from ..config import settings
        agent_name = settings.agent.name or "小助手"
        self.chat_title.setText(f"智能助理-{agent_name}")
        self.chat_title.setStyleSheet("""
            color: #202123;
            font-size: 16px;
            font-weight: 600;
        """)
        layout.addWidget(self.chat_title)
        layout.addStretch()

        return header

    def _create_input_area(self) -> "QWidget":
        """Create input area with file drop support"""
        container = self.QWidget()
        container.setStyleSheet("""
            QWidget {
                background-color: #ffffff;
                border-top: 1px solid #e5e5e5;
            }
        """)
        layout = self.QVBoxLayout(container)
        layout.setContentsMargins(24, 16, 24, 20)
        layout.setSpacing(12)

        # 文件预览区域
        self.files_widget = self.QWidget()
        self.files_widget.setStyleSheet("background: transparent;")
        self.files_layout = self.QHBoxLayout(self.files_widget)
        self.files_layout.setContentsMargins(0, 0, 0, 0)
        self.files_layout.setSpacing(8)
        self.files_widget.hide()
        layout.addWidget(self.files_widget)

        # 拖拽区域
        self.drop_area = self.QWidget()
        self.drop_area.setAcceptDrops(True)
        self.drop_area.setStyleSheet("""
            QWidget {
                background-color: #f7f7f8;
                border: 2px dashed #d9d9e3;
                border-radius: 12px;
            }
        """)
        drop_layout = self.QVBoxLayout(self.drop_area)
        drop_layout.setContentsMargins(16, 16, 16, 16)
        
        drop_label = self.QLabel("📁 拖拽文件或文件夹到此处，或输入消息")
        drop_label.setStyleSheet("color: #8e8ea0; font-size: 13px;")
        drop_label.setAlignment(self.Qt.AlignmentFlag.AlignCenter)
        drop_layout.addWidget(drop_label)
        
        layout.addWidget(self.drop_area)

        input_row = self.QHBoxLayout()
        input_row.setSpacing(12)

        self.input_field = self.HistoryLineEdit(agents_list=self.AGENTS_LIST)
        self.input_field.setPlaceholderText("与 AI 聊天，直接对话无需意图解析")
        self.input_field.setMinimumHeight(48)
        self.input_field.setStyleSheet("""
            QLineEdit {
                background-color: #f7f7f8;
                border: 1px solid #d9d9e3;
                border-radius: 12px;
                padding: 12px 16px;
                font-size: 14px;
                color: #202123;
            }
            QLineEdit:focus {
                border: 2px solid #10a37f;
                background-color: #ffffff;
            }
            QLineEdit::placeholder {
                color: #8e8ea0;
            }
        """)
        self.input_field.returnPressed.connect(self._on_send)
        self.input_field.textChanged.connect(self.input_field.on_text_changed)
        input_row.addWidget(self.input_field, 1)

        self.browse_button = self.QPushButton("📁")
        self.browse_button.setFixedSize(48, 48)
        self.browse_button.setToolTip("选择文件或文件夹")
        self.browse_button.setStyleSheet("""
            QPushButton {
                background-color: #f7f7f8;
                color: #202123;
                border: 1px solid #d9d9e3;
                border-radius: 12px;
                font-size: 20px;
            }
            QPushButton:hover {
                background-color: #e8f5e9;
                border: 1px solid #10a37f;
            }
            QPushButton:pressed {
                background-color: #c8e6c9;
            }
        """)
        self.browse_button.clicked.connect(self._on_browse)
        input_row.addWidget(self.browse_button)

        self.voice_button = create_voice_button(self)
        self.voice_button.setFixedSize(48, 48)
        input_row.addWidget(self.voice_button)
        self._voice_recording = False

        self.send_button = self.QPushButton("发送")
        self.send_button.setFixedSize(80, 48)
        self.send_button.setStyleSheet("""
            QPushButton {
                background-color: #10a37f;
                color: white;
                border: none;
                border-radius: 12px;
                font-size: 14px;
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: #0d8a6a;
            }
            QPushButton:pressed {
                background-color: #0a7259;
            }
            QPushButton:disabled {
                background-color: #d9d9e3;
                color: #8e8ea0;
            }
        """)
        self.send_button.clicked.connect(self._on_send)
        input_row.addWidget(self.send_button)

        layout.addLayout(input_row)

        hint = self.QLabel("按 Enter 发送，Shift+Enter 换行 | 支持拖拽文件")
        hint.setStyleSheet("color: #8e8ea0; font-size: 11px;")
        hint.setAlignment(self.Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(hint)

        token_layout = self.QHBoxLayout()
        token_layout.addStretch()
        
        self.token_label = self.QLabel("📊 Token: 0")
        self.token_label.setStyleSheet("color: #8e8ea0; font-size: 12px; padding: 4px;")
        self.token_label.setMinimumHeight(28)
        token_layout.addWidget(self.token_label)
        
        self.token_reset_btn = self.QPushButton("重置")
        self.token_reset_btn.setFixedSize(50, 28)
        self.token_reset_btn.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                color: #8e8ea0;
                border: 1px solid #d9d9e3;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: #f0f0f0;
            }
        """)
        self.token_reset_btn.clicked.connect(self._reset_token_count)
        token_layout.addWidget(self.token_reset_btn)
        
        layout.addLayout(token_layout)
        
        self._total_tokens = 0

        # 存储拖入的文件
        self.dropped_files: List[str] = []

        # 聊天模式上下文历史
        self._chat_history: List[dict] = []
        self._max_chat_history = 20

        # 启用拖拽
        self.drop_area.setAcceptDrops(True)
        self.drop_area.dragEnterEvent = self._drag_enter_event
        self.drop_area.dropEvent = self._drop_event

        return container

    def _show_welcome(self):
        """Show welcome message with dynamic greeting"""
        if hasattr(self, 'greeting_manager') and self.greeting_manager and hasattr(self, 'append_message'):
            from PyQt6.QtCore import QThread, pyqtSignal

            class GreetingThread(QThread):
                greeting_ready = pyqtSignal(str)

                def __init__(self, greeting_manager, parent_window):
                    super().__init__()
                    self.greeting_manager = greeting_manager
                    self.parent_window = parent_window

                def run(self):
                    try:
                        import asyncio
                        from ..multi_agent_system import get_global_loop

                        loop = get_global_loop()
                        future = asyncio.run_coroutine_threadsafe(
                            self.greeting_manager.get_greeting(),
                            loop
                        )
                        greeting = future.result(timeout=30)
                        self.greeting_ready.emit(greeting)
                    except Exception as e:
                        logger.debug(f"Greeting thread error: {e}")
                        from ..config import settings
                        user_name = settings.user.name or "主人"
                        agent_name = settings.agent.name or "智能助手"
                        self.greeting_ready.emit(f"👋 您好，{user_name}！我是{agent_name}。")

            def on_greeting(greeting):
                if hasattr(self, 'append_message'):
                    self.append_message("agent", greeting)

            self._greeting_thread = GreetingThread(self.greeting_manager, self)
            self._greeting_thread.greeting_ready.connect(on_greeting)
            self._greeting_thread.finished.connect(lambda: setattr(self, '_greeting_thread', None))
            self._greeting_thread.start()
        elif hasattr(self, 'append_message'):
            from ..config import settings
            user_name = settings.user.name or "主人"
            agent_name = settings.agent.name or "智能助手"
            welcome = f"""👋 您好，{user_name}！我是{agent_name}。

我可以帮您完成以下任务：

**🔍 网络搜索** - 搜索互联网上的信息
**🌤️ 天气预报** - 查询任意城市的天气
**📧 邮件助手** - 发送和管理邮件
**🎵 音乐播放** - 播放音乐和音频

有什么我可以帮您的吗？"""
            self.append_message("agent", welcome)

    def _clear_chat(self):
        """Clear chat history"""
        self.messages_list.clear()
        self._show_welcome()

    def _init_conversation_manager(self):
        """Initialize conversation manager"""
        from .conversation_manager import conversation_manager
        self.conv_manager = conversation_manager

    def _load_initial_messages(self):
        """Load initial messages for current conversation"""
        conv = self.conv_manager.get_conversation()
        if conv and conv.messages:
            self._load_conversation_messages()
        else:
            self._show_welcome()

    def _clear_conversation(self):
        """Clear conversation display (keeps history for LLM context)"""
        if not hasattr(self, 'conv_manager'):
            self._init_conversation_manager()
        
        self.conv_manager.clear_messages()
        self.messages_list.clear()
        
        from ..config import settings
        user_name = settings.user.name or "主人"
        agent_name = settings.agent.name or "智能助手"
        simple_welcome = f"👋 {user_name}，对话已清空。有什么我可以帮您的吗？"
        self._append_message_without_save("agent", simple_welcome)
        logger.info("🗑️ 对话显示已清空（历史记录保留）")

    def _load_conversation_messages(self):
        """Load messages for current conversation"""
        self.messages_list.clear()
        
        conv = self.conv_manager.get_conversation()
        if conv:
            for msg in conv.messages:
                self._append_message_without_save(msg.role, msg.content)
            
            if not conv.messages:
                self._show_welcome()

    def append_message(self, role: str, content: str, metadata: dict = None):
        """Add message to chat and save to conversation"""
        self._append_message_without_save(role, content, metadata)
        if hasattr(self, 'conv_manager') and role != "system":
            self.conv_manager.add_message(role, content, metadata)
            
            from ..memory.history_manager import history_manager
            history_manager.add_message(role, content)
        
        if role == "agent" and hasattr(self, '_pending_auto_speak') and self._pending_auto_speak:
            text_to_speak = self._pending_auto_speak
            self._pending_auto_speak = None
            self._start_auto_speak_after_message_added(text_to_speak)

    def _append_message_without_save(self, role: str, content: str, metadata: dict = None):
        """Add message to chat without saving to conversation"""
        from PyQt6.QtWidgets import QListWidgetItem, QHBoxLayout, QVBoxLayout, QLabel
        
        timestamp = datetime.now().strftime("%H:%M")
        
        item = self.QListWidgetItem()
        item.setFlags(item.flags() & ~self.Qt.ItemFlag.ItemIsSelectable)
        
        msg_widget = self.QWidget()
        msg_widget.setStyleSheet("background: transparent;")
        msg_widget._role = role
        msg_widget._content = content
        msg_widget._audio_cache = None
        
        # 设置 skip_auto_speak 标志，用于控制播放按钮的显示
        skip_auto_speak = metadata and metadata.get('skip_auto_speak', False)
        msg_widget._skip_auto_speak = skip_auto_speak
        
        from loguru import logger
        logger.info(f"🔍 _append_message_without_save: role={role}, metadata={metadata}, skip_auto_speak={skip_auto_speak}")
        
        outer_layout = self.QVBoxLayout(msg_widget)
        outer_layout.setContentsMargins(24, 8, 24, 8)
        outer_layout.setSpacing(4)
        
        if role == "user":
            container = self.QWidget()
            container.setStyleSheet("background: transparent;")
            container_layout = self.QHBoxLayout(container)
            container_layout.setContentsMargins(0, 0, 0, 0)
            container_layout.addStretch()
            
            bubble = self.BubbleLabel(content, list_widget=self.messages_list)
            bubble.setWordWrap(True)
            bubble.setTextInteractionFlags(self.Qt.TextInteractionFlag.TextSelectableByMouse)
            bubble.setStyleSheet("""
                QLabel {
                    background-color: #10a37f;
                    color: white;
                    padding: 12px 16px;
                    border-radius: 16px;
                    font-size: 14px;
                    line-height: 1.5;
                }
            """)
            list_width = self.messages_list.width() if hasattr(self, 'messages_list') else 600
            bubble.setMaximumWidth(max(300, list_width - 150))
            container_layout.addWidget(bubble)
            
            time_label = self.QLabel(timestamp)
            time_label.setStyleSheet("color: #8e8ea0; font-size: 11px;")
            time_label.setAlignment(self.Qt.AlignmentFlag.AlignRight)
            
            outer_layout.addWidget(container)
            outer_layout.addWidget(time_label)
        elif role == "system":
            bubble = self.QLabel(content)
            bubble.setWordWrap(True)
            bubble.setStyleSheet("""
                QLabel {
                    background-color: #fff3cd;
                    color: #856404;
                    padding: 8px 12px;
                    border-radius: 8px;
                    font-size: 12px;
                    font-style: italic;
                }
            """)
            bubble.setMaximumWidth(400)
            bubble.setAlignment(self.Qt.AlignmentFlag.AlignCenter)
            
            container = self.QWidget()
            container.setStyleSheet("background: transparent;")
            container_layout = self.QHBoxLayout(container)
            container_layout.setContentsMargins(0, 0, 0, 0)
            container_layout.addStretch()
            container_layout.addWidget(bubble)
            container_layout.addStretch()
            
            outer_layout.addWidget(container)
        else:
            container = self.QWidget()
            container.setStyleSheet("background: transparent;")
            container_layout = self.QHBoxLayout(container)
            container_layout.setContentsMargins(0, 0, 0, 0)
            container_layout.setSpacing(12)
            
            avatar = self._create_avatar_label()
            container_layout.addWidget(avatar)
            
            content_widget = self.QWidget()
            content_widget.setStyleSheet("background: transparent;")
            content_layout = self.QVBoxLayout(content_widget)
            content_layout.setContentsMargins(0, 0, 0, 0)
            content_layout.setSpacing(4)
            
            bubble_row = self.QWidget()
            bubble_row.setStyleSheet("background: transparent;")
            bubble_layout = self.QHBoxLayout(bubble_row)
            bubble_layout.setContentsMargins(0, 0, 0, 0)
            bubble_layout.setSpacing(8)
            
            rendered = self._render_markdown(content)
            bubble = self.BubbleLabel(rendered, list_widget=self.messages_list)
            bubble.setWordWrap(True)
            bubble.setTextInteractionFlags(
                self.Qt.TextInteractionFlag.TextSelectableByMouse | 
                self.Qt.TextInteractionFlag.LinksAccessibleByMouse
            )
            bubble.setOpenExternalLinks(True)
            bubble.setStyleSheet("""
                QLabel {
                    background-color: #f7f7f8;
                    color: #202123;
                    padding: 12px 16px;
                    border-radius: 12px;
                    font-size: 14px;
                    line-height: 1.6;
                }
            """)
            list_width = self.messages_list.width() if hasattr(self, 'messages_list') else 600
            max_width = max(300, list_width - 200)
            bubble.setMaximumWidth(max_width)
            bubble_layout.addWidget(bubble, 1)
            
            skip_auto_speak_check = metadata and metadata.get('skip_auto_speak', False)
            from ..tts import get_tts_manager
            from loguru import logger
            logger.info(f"🔍 _append_message_without_save (play button): metadata={metadata}, skip_auto_speak_check={skip_auto_speak_check}, tts_enabled={get_tts_manager().is_enabled()}")
            
            tts = get_tts_manager()
            if tts.is_enabled() and not skip_auto_speak_check:
                play_btn = self.QPushButton("▶")
                play_btn.setFixedSize(28, 28)
                play_btn.setToolTip("播放语音")
                play_btn.setObjectName("tts_play_btn")
                play_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #10a37f;
                        color: white;
                        border: none;
                        border-radius: 14px;
                        font-size: 12px;
                    }
                    QPushButton:hover {
                        background-color: #0d8a6a;
                    }
                """)
                msg_id = id(msg_widget)
                play_btn.clicked.connect(lambda checked, mid=msg_id, txt=content, btn=play_btn: self._toggle_tts(mid, txt, btn))
                bubble_layout.addWidget(play_btn)
            else:
                play_btn = self.QPushButton("▶")
                play_btn.setFixedSize(28, 28)
                play_btn.setToolTip("播放语音")
                play_btn.setObjectName("tts_play_btn")
                play_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #10a37f;
                        color: white;
                        border: none;
                        border-radius: 14px;
                        font-size: 12px;
                    }
                    QPushButton:hover {
                        background-color: #0d8a6a;
                    }
                """)
                play_btn.setVisible(False)
                msg_id = id(msg_widget)
                play_btn.clicked.connect(lambda checked, mid=msg_id, txt=content, btn=play_btn: self._toggle_tts(mid, txt, btn))
                bubble_layout.addWidget(play_btn)
                logger.info(f"🔍 _append_message_without_save: play button hidden (skip_auto_speak_check={skip_auto_speak_check})")
            
            content_layout.addWidget(bubble_row)
            
            time_label = self.QLabel(timestamp)
            time_label.setStyleSheet("color: #8e8ea0; font-size: 11px;")
            content_layout.addWidget(time_label)
            
            container_layout.addWidget(content_widget, 1)
            container_layout.addStretch()
            
            outer_layout.addWidget(container)
        
        msg_widget.adjustSize()
        msg_widget.updateGeometry()
        
        total_height = msg_widget.sizeHint().height()
        
        if total_height < 50:
            total_height = 50
        
        item.setSizeHint(self.QSize(self.messages_list.width(), total_height + 16))
        
        self.messages_list.addItem(item)
        self.messages_list.setItemWidget(item, msg_widget)
        
        from PyQt6.QtCore import QTimer
        QTimer.singleShot(100, self.smooth_scroll_to_bottom)
    
    def _reset_play_button(self):
        """重置播放按钮状态"""
        if hasattr(self, '_current_playing_btn') and self._current_playing_btn:
            try:
                self._current_playing_btn.setText("▶")
                self._current_playing_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #10a37f;
                        color: white;
                        border: none;
                        border-radius: 14px;
                        font-size: 12px;
                    }
                    QPushButton:hover {
                        background-color: #0d8a6a;
                    }
                """)
            except:
                pass
        self._current_playing_msg = None
        self._current_playing_btn = None
        self._pending_auto_speak = None
    
    def _toggle_tts(self, msg_id: int, text: str, btn):
        """切换播放/停止语音"""
        import re
        from ..tts import get_tts_manager
        
        tts = get_tts_manager()
        
        if hasattr(self, '_current_playing_msg') and self._current_playing_msg == msg_id:
            tts.stop()
            self._reset_play_button()
        else:
            self._reset_play_button()
            
            clean_text = re.sub(r'[📍🔑📱🌐✅❌📋🎯📎📝🔊]', '', text)
            clean_text = re.sub(r'https?://\S+', '', clean_text)
            clean_text = clean_text.strip()
            
            if clean_text and tts.is_enabled():
                self._current_playing_msg = msg_id
                self._current_playing_btn = btn
                btn.setText("⏹")
                btn.setStyleSheet("""
                    QPushButton {
                        background-color: #ef4444;
                        color: white;
                        border: none;
                        border-radius: 14px;
                        font-size: 12px;
                    }
                    QPushButton:hover {
                        background-color: #dc2626;
                    }
                """)
                
                msg_widget = None
                for i in range(self.messages_list.count()):
                    item = self.messages_list.item(i)
                    widget = self.messages_list.itemWidget(item)
                    if widget and id(widget) == msg_id:
                        msg_widget = widget
                        break
                
                import threading
                import os
                
                def play_and_reset():
                    audio_path = None
                    if msg_widget and hasattr(msg_widget, '_audio_cache') and msg_widget._audio_cache:
                        audio_path = msg_widget._audio_cache
                        if not os.path.exists(audio_path):
                            audio_path = None
                    
                    if audio_path:
                        tts.play_cached(audio_path)
                    else:
                        audio_path = tts.speak_sync_with_cache(clean_text)
                        if audio_path and msg_widget:
                            msg_widget._audio_cache = audio_path
                    
                    self.signal_helper.emit_reset_play_button()
                
                thread = threading.Thread(target=play_and_reset, daemon=True)
                thread.start()

    def show_progress(self, message: str, progress: int = -1):
        """Show progress message"""
        from PyQt6.QtWidgets import QListWidgetItem, QHBoxLayout, QVBoxLayout, QLabel
        
        timestamp = datetime.now().strftime("%H:%M")
        
        item = self.QListWidgetItem()
        item.setFlags(item.flags() & ~self.Qt.ItemFlag.ItemIsSelectable)
        
        msg_widget = self.QWidget()
        msg_widget.setStyleSheet("background: transparent;")
        outer_layout = self.QVBoxLayout(msg_widget)
        outer_layout.setContentsMargins(24, 8, 24, 8)
        outer_layout.setSpacing(4)
        
        container = self.QWidget()
        container.setStyleSheet("background: transparent;")
        container_layout = self.QHBoxLayout(container)
        container_layout.setContentsMargins(0, 0, 0, 0)
        container_layout.setSpacing(12)
        
        avatar = self._create_avatar_label()
        container_layout.addWidget(avatar)
        
        if progress >= 0:
            filled = int(progress / 10)
            empty = 10 - filled
            bar = "█" * filled + "░" * empty
            content = f"{message}\n[{bar}] {progress}%"
        else:
            content = f"⏳ {message}..."
        
        bubble = self.QLabel(content)
        bubble.setWordWrap(True)
        bubble.setStyleSheet("""
            QLabel {
                background-color: #fff8e6;
                color: #7c5800;
                padding: 12px 16px;
                border-radius: 12px;
                font-size: 13px;
                border: 1px solid #ffd966;
            }
        """)
        container_layout.addWidget(bubble)
        container_layout.addStretch()
        
        outer_layout.addWidget(container)
        
        time_label = self.QLabel(timestamp)
        time_label.setStyleSheet("color: #8e8ea0; font-size: 11px;")
        outer_layout.addWidget(time_label)
        
        msg_widget.adjustSize()
        height = msg_widget.sizeHint().height()
        
        item.setSizeHint(self.QSize(self.messages_list.width(), height + 8))
        
        self.messages_list.addItem(item)
        self.messages_list.setItemWidget(item, msg_widget)
        self.smooth_scroll_to_bottom()
        
        from PyQt6.QtWidgets import QApplication
        QApplication.processEvents()

    def _create_avatar_label(self):
        """Create avatar label from settings"""
        from ..config import settings
        from pathlib import Path
        
        avatar_text = settings.agent.avatar or "🤖"
        
        avatar = self.QLabel()
        avatar.setFixedSize(36, 36)
        avatar.setAlignment(self.Qt.AlignmentFlag.AlignCenter)
        
        if avatar_text.startswith("/") or avatar_text.startswith("\\") or ":" in avatar_text:
            avatar_path = Path(avatar_text)
            if avatar_path.exists():
                from PyQt6.QtGui import QPixmap
                pixmap = QPixmap(str(avatar_path))
                if not pixmap.isNull():
                    scaled = pixmap.scaled(36, 36, self.Qt.AspectRatioMode.KeepAspectRatioByExpanding, self.Qt.TransformationMode.SmoothTransformation)
                    avatar.setPixmap(scaled)
                    avatar.setStyleSheet("border-radius: 8px;")
                    return avatar
        
        avatar.setText(avatar_text)
        avatar.setStyleSheet("""
            font-size: 24px;
            background-color: #10a37f;
            border-radius: 8px;
            padding: 4px;
        """)
        
        return avatar

    def _render_markdown(self, text: str) -> str:
        """Convert markdown to HTML"""
        try:
            html = self.markdown.markdown(text, extensions=['fenced_code', 'tables', 'nl2br'])
            html = html.replace('<p>', '').replace('</p>', '<br>')
            html = html.replace('<code>', '<span style="background-color: #e8e8e8; padding: 2px 6px; border-radius: 4px; font-family: Consolas, monospace;">')
            html = html.replace('</code>', '</span>')
            html = html.replace('<pre>', '<div style="background-color: #f7f7f8; padding: 12px; border-radius: 8px; margin: 8px 0; overflow-x: auto;"><code>')
            html = html.replace('</pre>', '</code></div>')
            html = html.replace('<strong>', '<b>')
            html = html.replace('</strong>', '</b>')
            html = html.replace('<ul>', '')
            html = html.replace('</ul>', '')
            html = html.replace('<li>', '&nbsp;&nbsp;• ')
            html = html.replace('</li>', '<br>')
            html = html.replace('<ol>', '')
            html = html.replace('</ol>', '')
            import re
            html = re.sub(r'<li>(.*?)</li>', r'&nbsp;&nbsp;\1<br>', html)
            
            url_pattern = r'(?<!href=")(?<!src=")(https?://[^\s<>\[\]"\'\u4e00-\u9fff]+)'
            
            def make_url_link(match):
                url = match.group(1)
                clean_url = url.rstrip('.,;:!?）】」』"\'')
                return f'<a href="{clean_url}" style="color: #10a37f; text-decoration: underline;">{clean_url}</a>'
            
            html = re.sub(url_pattern, make_url_link, html)
            
            file_path_pattern = r'([A-Za-z]:\\[^\s<>\[\]"\'\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]*[^\s<>\[\]"\'。，，；：！？、）】」』"]*)'
            
            def make_file_link(match):
                path = match.group(1)
                if len(path) < 5:
                    return path
                import os
                if os.path.exists(path):
                    return f'<a href="file:///{path}" style="color: #10a37f; text-decoration: underline;">{path}</a>'
                return path
            
            html = re.sub(file_path_pattern, make_file_link, html)
            
            return html
        except Exception:
            return text.replace('\n', '<br>')

    def _show_web_server_dialog(self, url: str, qr_code: str):
        """显示 Web 服务器信息对话框（包含二维码）"""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton, QHBoxLayout
        from PyQt6.QtGui import QPixmap, QImage
        from PyQt6.QtCore import Qt
        import base64
        
        logger.info(f"显示Web服务器对话框: url={url}, qr_code长度={len(qr_code) if qr_code else 0}")
        
        dialog = QDialog(self.window)
        dialog.setWindowTitle("Web 服务已启动")
        dialog.setFixedSize(400, 450)
        
        layout = QVBoxLayout(dialog)
        
        title = QLabel("📱 扫码访问智能助手")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("font-size: 18px; font-weight: bold; margin: 10px;")
        layout.addWidget(title)
        
        qr_label = None
        if qr_code and qr_code.startswith('data:image'):
            try:
                base64_data = qr_code.split(',')[1]
                image_data = base64.b64decode(base64_data)
                image = QImage()
                if image.loadFromData(image_data):
                    pixmap = QPixmap.fromImage(image)
                    
                    qr_label = QLabel()
                    qr_label.setPixmap(pixmap.scaled(200, 200, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
                    qr_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    layout.addWidget(qr_label)
                    logger.info("二维码图片加载成功")
                else:
                    logger.error("二维码图片加载失败")
            except Exception as e:
                logger.error(f"解析二维码失败: {e}")
        
        if qr_label is None:
            no_qr = QLabel("⚠️ 二维码生成失败")
            no_qr.setAlignment(Qt.AlignmentFlag.AlignCenter)
            no_qr.setStyleSheet("font-size: 14px; color: #e74c3c; margin: 10px;")
            layout.addWidget(no_qr)
        
        url_label = QLabel(f"🌐 地址: {url}")
        url_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        url_label.setStyleSheet("font-size: 14px; margin: 10px;")
        url_label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        layout.addWidget(url_label)
        
        hint = QLabel("请确保手机和电脑在同一 WiFi 网络下")
        hint.setAlignment(Qt.AlignmentFlag.AlignCenter)
        hint.setStyleSheet("font-size: 12px; color: #666; margin: 10px;")
        layout.addWidget(hint)
        
        close_btn = QPushButton("关闭")
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: #667eea;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #5a6fd6;
            }
        """)
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)
        
        dialog.exec()

    def _get_web_html(self) -> str:
        """获取 Web 界面 HTML"""
        return """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>智能助手</title>
    <style>
        * { box-sizing: border-box; margin: 0; padding: 0; }
        body { 
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Helvetica, Arial, sans-serif; 
            max-width: 800px; 
            margin: 0 auto; 
            padding: 16px; 
            background: #f5f5f5;
            min-height: 100vh;
        }
        h1 { text-align: center; color: #333; margin-bottom: 16px; font-size: 1.5em; }
        #chat { 
            height: calc(100vh - 160px); 
            min-height: 300px;
            overflow-y: auto; 
            border: 1px solid #ddd; 
            padding: 16px; 
            margin-bottom: 16px; 
            border-radius: 12px; 
            background: #fff;
        }
        .message { 
            margin: 12px 0; 
            padding: 12px 16px; 
            border-radius: 12px; 
            max-width: 85%;
            word-wrap: break-word;
        }
        .user { background: #e3f2fd; margin-left: auto; text-align: right; }
        .agent { background: #e8f5e9; margin-right: auto; }
        .label { font-size: 0.8em; color: #666; margin-bottom: 4px; }
        #input-area { display: flex; gap: 8px; }
        #input { 
            flex: 1; 
            padding: 12px; 
            border: 1px solid #ddd; 
            border-radius: 8px; 
            font-size: 16px;
        }
        #send { 
            padding: 12px 24px; 
            background: #4CAF50; 
            color: white; 
            border: none; 
            border-radius: 8px; 
            cursor: pointer; 
            font-size: 16px;
        }
        #send:hover { background: #45a049; }
        #send:disabled { background: #ccc; cursor: not-allowed; }
        .loading { color: #666; font-style: italic; }
    </style>
</head>
<body>
    <h1>🤖 智能助手</h1>
    <div id="chat"></div>
    <div id="input-area">
        <input type="text" id="input" placeholder="输入消息..." autocomplete="off">
        <button id="send">发送</button>
    </div>
    <script>
        const chat = document.getElementById('chat');
        const input = document.getElementById('input');
        const send = document.getElementById('send');

        function addMessage(role, content) {
            const div = document.createElement('div');
            div.className = 'message ' + role;
            div.innerHTML = '<div class="label">' + (role === 'user' ? '👤 你' : '🤖 助手') + '</div>' +
                           '<div>' + content.replace(/\\n/g, '<br>') + '</div>';
            chat.appendChild(div);
            chat.scrollTop = chat.scrollHeight;
        }

        function setLoading(show) {
            if (show) {
                const div = document.createElement('div');
                div.className = 'message agent loading';
                div.id = 'loading';
                div.textContent = '思考中...';
                chat.appendChild(div);
                chat.scrollTop = chat.scrollHeight;
            } else {
                const loading = document.getElementById('loading');
                if (loading) loading.remove();
            }
        }

        async function sendMessage() {
            const message = input.value.trim();
            if (!message) return;
            
            addMessage('user', message);
            input.value = '';
            input.disabled = true;
            send.disabled = true;
            setLoading(true);
            
            try {
                const response = await fetch('/chat', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({message: message})
                });
                const data = await response.json();
                setLoading(false);
                addMessage('agent', data.response);
            } catch (e) {
                setLoading(false);
                addMessage('agent', '错误: ' + e.message);
            } finally {
                input.disabled = false;
                send.disabled = false;
                input.focus();
            }
        }

        send.onclick = sendMessage;
        input.onkeypress = (e) => { if (e.key === 'Enter') sendMessage(); };
        input.focus();
    </script>
</body>
</html>
"""

    def _show_settings_dialog(self):
        """Show settings dialog"""
        try:
            from .settings_dialog import SettingsDialog
            dialog = SettingsDialog(self.window)
            dialog.settings_changed.connect(self._on_settings_changed)
            dialog.exec()
        except Exception as e:
            logger.error(f"Failed to open settings dialog: {e}")
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.information(self.window, "设置", f"设置功能暂时不可用\n\n错误: {e}")
    
    def _show_music_minimized(self):
        """显示音乐播放器最小化条目"""
        if hasattr(self, 'music_minimized_item'):
            self.music_minimized_item.show()
    
    def _show_music_player(self):
        """Show music player window"""
        try:
            from ..music import MusicPlayerWidget
            from PyQt6.QtWidgets import QMainWindow
            from PyQt6.QtCore import Qt
            
            if not hasattr(self, '_music_window') or self._music_window is None:
                self._music_window = QMainWindow()
                self._music_window.setWindowTitle("🎵 音乐播放器")
                self._music_window.setMinimumSize(900, 600)
                
                self._music_window.setWindowFlags(
                    Qt.WindowType.Window |
                    Qt.WindowType.WindowTitleHint |
                    Qt.WindowType.WindowCloseButtonHint
                )
                self._music_window.setWindowFlag(Qt.WindowType.WindowMinimizeButtonHint, False)
                self._music_window.setWindowFlag(Qt.WindowType.WindowMaximizeButtonHint, False)
                
                shared_player = self._get_shared_music_player()
                self._music_widget = MusicPlayerWidget(player=shared_player)
                self._music_window.setCentralWidget(self._music_widget)
                
                self._music_widget.song_changed.connect(self._on_music_song_changed)
                self._music_widget.state_changed.connect(self._on_music_state_changed)
                
                self._music_window.closeEvent = self._on_music_window_close
            
            self._music_window.showNormal()
            self._music_window.show()
            self._music_window.raise_()
            self._music_window.activateWindow()
            
        except Exception as e:
            logger.error(f"Failed to open music player: {e}")
            import traceback
            traceback.print_exc()
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.information(self.window, "音乐播放器", f"音乐播放器暂时不可用\n\n错误: {e}")
    
    def _get_shared_music_player(self):
        """获取共享的音乐播放器实例"""
        try:
            from ..multi_agent_system import multi_agent_system
            if multi_agent_system and multi_agent_system.master:
                music_agent = multi_agent_system.master.sub_agents.get('music_agent')
                if music_agent and hasattr(music_agent, 'player') and music_agent.player:
                    logger.info("🎵 使用 music_agent 的共享播放器实例")
                    return music_agent.player
        except Exception as e:
            logger.debug(f"获取 music_agent 播放器失败: {e}")
        
        from ..config import settings
        from ..music import MusicPlayer
        music_library = settings.directory.get_music_library()
        logger.info("🎵 创建新的音乐播放器实例")
        return MusicPlayer(music_library=music_library)
    
    def _on_music_window_close(self, event):
        """Handle music player window close - hide window but keep playing"""
        try:
            # 只隐藏窗口，不停止音乐
            if hasattr(self, '_music_window') and self._music_window:
                self._music_window.hide()
            logger.info("🎵 音乐播放器窗口已隐藏")
        except Exception as e:
            logger.error(f"隐藏音乐播放器窗口失败: {e}")
        event.ignore()  # 忽略关闭事件，改为隐藏
    
    def _stop_music_and_hide(self):
        """停止音乐并隐藏最小化条目"""
        try:
            stopped = False
            shared_player = self._get_shared_music_player()
            if shared_player and shared_player.is_playing:
                shared_player.stop()
                stopped = True
                logger.info("🎵 音乐已停止（通过共享播放器）")
            
            if not stopped and hasattr(self, '_music_widget') and self._music_widget:
                self._music_widget.player.stop()
                stopped = True
                logger.info("🎵 音乐已停止（通过 widget）")
            
            if hasattr(self, 'music_minimized_item'):
                self.music_minimized_item.hide()
                
        except Exception as e:
            logger.error(f"停止音乐失败: {e}")
    
    def _on_music_song_changed(self, song_data):
        """音乐播放器歌曲变化"""
        title = song_data.get("title", "未知歌曲")
        self.append_message("agent", f"🎵 正在播放: {title}")
        
        # 更新最小化条目状态
        if hasattr(self, 'music_minimized_status'):
            self.music_minimized_status.setText(f"正在播放: {title[:15]}..." if len(title) > 15 else f"正在播放: {title}")
        
        # 显示最小化条目
        if hasattr(self, 'music_minimized_item'):
            self.music_minimized_item.show()
    
    def _on_music_state_changed(self, state):
        """音乐播放器状态变化"""
        logger.info(f"🎵 播放器状态: {state}")
        
        if hasattr(self, 'music_minimized_status'):
            if state.get("playing", False):
                self.music_minimized_status.setText("正在播放...")
            else:
                self.music_minimized_status.setText("已暂停")
    
    def _show_homeassistant_dashboard(self):
        """显示智能家居控制面板"""
        try:
            from PyQt6.QtWidgets import QMainWindow
            from PyQt6.QtCore import Qt
            from .homeassistant_dashboard import HomeAssistantDashboard
            
            if not hasattr(self, '_ha_window') or self._ha_window is None:
                self._ha_window = QMainWindow()
                self._ha_window.setWindowTitle("🏠 智能家居控制")
                self._ha_window.setMinimumSize(600, 500)
                
                self._ha_window.setWindowFlags(
                    Qt.WindowType.Window |
                    Qt.WindowType.WindowTitleHint |
                    Qt.WindowType.WindowCloseButtonHint
                )
                
                self._ha_dashboard = HomeAssistantDashboard()
                self._ha_window.setCentralWidget(self._ha_dashboard)
                
                self._ha_dashboard.refresh_requested.connect(self._refresh_ha_entities)
                self._ha_dashboard.control_requested.connect(self._control_ha_device)
                
                self._ha_window.closeEvent = self._on_ha_window_close
            
            self._ha_window.showNormal()
            self._ha_window.show()
            self._ha_window.raise_()
            self._ha_window.activateWindow()
            
            self._refresh_ha_entities()
            
        except Exception as e:
            logger.error(f"打开智能家居控制面板失败: {e}")
            import traceback
            traceback.print_exc()
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.information(self.window, "智能家居", f"智能家居控制面板暂时不可用\n\n错误: {e}")
    
    def _on_ha_window_close(self, event):
        """智能家居窗口关闭事件"""
        try:
            if hasattr(self, '_ha_window') and self._ha_window:
                self._ha_window.hide()
            logger.info("🏠 智能家居控制面板已隐藏")
        except Exception as e:
            logger.error(f"隐藏智能家居控制面板失败: {e}")
        event.ignore()
    
    async def _get_ha_agent(self):
        """获取 Home Assistant 智能体"""
        try:
            from ..multi_agent_system import multi_agent_system
            if multi_agent_system and multi_agent_system.master:
                return multi_agent_system.master.sub_agents.get('homeassistant_agent')
        except Exception as e:
            logger.debug(f"获取 homeassistant_agent 失败: {e}")
        return None
    
    def _refresh_ha_entities(self):
        """刷新智能家居设备列表"""
        async def _refresh():
            ha_agent = await self._get_ha_agent()
            if ha_agent and ha_agent.api:
                try:
                    await ha_agent._refresh_entity_cache()
                    entities = ha_agent._entity_cache
                    
                    from PyQt6.QtWidgets import QApplication
                    app = QApplication.instance()
                    if app:
                        def update_ui():
                            if hasattr(self, '_ha_dashboard'):
                                self._ha_dashboard.set_entities(entities)
                        app.invokeLater(update_ui)
                        
                except Exception as e:
                    logger.error(f"刷新设备列表失败: {e}")
            else:
                if hasattr(self, '_ha_dashboard'):
                    self._ha_dashboard.set_status("未配置 Home Assistant")
        
        import asyncio
        try:
            loop = asyncio.get_event_loop()
            if loop.is_running():
                asyncio.create_task(_refresh())
            else:
                loop.run_until_complete(_refresh())
        except Exception as e:
            logger.error(f"刷新设备列表失败: {e}")
    
    def _control_ha_device(self, entity_id: str, action: str, params: dict):
        """控制智能家居设备"""
        async def _control():
            ha_agent = await self._get_ha_agent()
            if ha_agent and ha_agent.api:
                try:
                    if action == 'on':
                        await ha_agent.api.turn_on(entity_id)
                    elif action == 'off':
                        await ha_agent.api.turn_off(entity_id)
                    else:
                        await ha_agent.api.toggle(entity_id)
                    
                    self._refresh_ha_entities()
                    
                except Exception as e:
                    logger.error(f"控制设备失败: {e}")
        
        import asyncio
        try:
            loop = asyncio.get_event_loop()
            if loop.is_running():
                asyncio.create_task(_control())
            else:
                loop.run_until_complete(_control())
        except Exception as e:
            logger.error(f"控制设备失败: {e}")
    
    def _on_settings_changed(self):
        """Handle settings changed"""
        from ..config import settings
        settings.reload()
        
        if hasattr(self, '_music_widget') and self._music_widget:
            new_path = settings.directory.get_music_library()
            if self._music_widget.player.music_library != new_path:
                self._music_widget.player.music_library = new_path
                self._music_widget.library_path_label.setText(f"📁 {new_path}")
                self._music_widget._scan_library()
        
        agent_name = settings.agent.name or "小助手"
        self.chat_title.setText(f"智能助理-{agent_name}")
        
        # 重新初始化 TTSManager 以应用新的设置
        from ..tts import init_tts, get_tts_manager
        init_tts()  # 重新初始化 TTSManager
        tts_enabled = get_tts_manager().is_enabled()
        self._update_play_buttons_visibility(tts_enabled)
        
        if hasattr(self, 'append_message'):
            self.append_message("agent", "✅ 设置已保存，头像设置将在下次发送消息时生效。")
    
    def _update_play_buttons_visibility(self, visible: bool):
        """更新所有播放按钮的可见性"""
        if not hasattr(self, 'messages_list'):
            return
        
        for i in range(self.messages_list.count()):
            item = self.messages_list.item(i)
            widget = self.messages_list.itemWidget(item)
            if widget:
                # 检查消息是否应该跳过自动语音合成
                skip_auto_speak = getattr(widget, '_skip_auto_speak', False)
                if skip_auto_speak:
                    continue
                
                play_btns = widget.findChildren(self.QPushButton, "tts_play_btn")
                for btn in play_btns:
                    btn.setVisible(visible)

    def _init_multi_agent_system(self):
        """Initialize multi-agent system in background"""
        import threading
        
        def init_in_background():
            try:
                from ..multi_agent_system import multi_agent_system, get_global_loop
                from ..agents.news_agent import NewsAgent
                from ..greeting.greeting_manager import GreetingManager
                from ..memory.long_term_memory import LongTermMemory
                from ..utils.token_counter import token_counter
                import asyncio

                loop = get_global_loop()

                async def init_system():
                    if not multi_agent_system._initialized:
                        await multi_agent_system.initialize()
                    
                    self.multi_agent = multi_agent_system

                    if self.multi_agent.master:
                        self.multi_agent.master.set_notification_callback(self._show_notification)
                    
                    self.multi_agent.set_auto_speak_callback(self.signal_helper.emit_auto_speak)
                    
                    token_counter.set_callback(self.signal_helper.emit_update_token)

                    news_agent = NewsAgent()
                    await news_agent.start()

                    memory = LongTermMemory()
                    greeting_manager = GreetingManager(
                        memory, 
                        None, 
                        news_agent, 
                        master_agent=self.multi_agent.master
                    )

                    self.news_agent = news_agent
                    self.greeting_manager = greeting_manager

                future = asyncio.run_coroutine_threadsafe(init_system(), loop)
                future.result(timeout=30)

                self.multi_agent.loop = loop

                logger.info("✅ GUI 初始化完成")
            except Exception as e:
                logger.warning(f"Failed to init multi-agent system: {e}")
                import traceback
                traceback.print_exc()
                self.multi_agent = None

        thread = threading.Thread(target=init_in_background, daemon=True)
        thread.start()

    def _on_notification_signal(self, message: str):
        """Handle notification signal from other threads (runs in main thread)"""
        self.append_message("agent", message)

    def _on_append_message_signal(self, role: str, content: str):
        """Handle append message signal from other threads (runs in main thread)"""
        self.append_message(role, content)

    def _on_auto_speak(self, text: str):
        """Handle auto speak signal - store text and wait for message to be added"""
        from ..tts import get_tts_manager
        
        tts = get_tts_manager()
        if not tts.is_enabled():
            return
        
        self._pending_auto_speak = text
    
    def _start_auto_speak_after_message_added(self, text: str):
        """Start auto speak after the agent message has been added to the list"""
        from ..tts import get_tts_manager
        import threading
        import re
        import time
        
        start_time = time.time()
        
        tts = get_tts_manager()
        if not tts.is_enabled():
            return
        
        self._reset_play_button()
        
        last_widget = None
        btn = None
        msg_id = None
        
        for i in range(self.messages_list.count() - 1, -1, -1):
            item = self.messages_list.item(i)
            widget = self.messages_list.itemWidget(item)
            if widget and hasattr(widget, '_role') and widget._role == "agent":
                last_widget = widget
                break
        
        if last_widget:
            play_btns = last_widget.findChildren(self.QPushButton, "tts_play_btn")
            if play_btns:
                btn = play_btns[-1]
                msg_id = id(last_widget)
        
        if last_widget and btn and msg_id:
            self._current_playing_msg = msg_id
            self._current_playing_btn = btn
            btn.setText("⏹")
            btn.setStyleSheet("""
                QPushButton {
                    background-color: #ef4444;
                    color: white;
                    border: none;
                    border-radius: 14px;
                    font-size: 12px;
                }
                QPushButton:hover {
                    background-color: #dc2626;
                }
            """)
            btn.repaint()
        
        clean_text = re.sub(r'[📍🔑📱🌐✅❌📋🎯📎📝🔊]', '', text)
        clean_text = re.sub(r'https?://\S+', '', clean_text)
        clean_text = clean_text.strip()
        
        from loguru import logger
        logger.info(f"🔊 GUI开始语音处理，文本长度: {len(clean_text)} 字符")
        
        def speak_in_thread():
            synth_start = time.time()
            audio_path = None
            if last_widget and hasattr(last_widget, '_audio_cache') and last_widget._audio_cache:
                import os
                if os.path.exists(last_widget._audio_cache):
                    audio_path = last_widget._audio_cache
                    logger.info(f"🔊 使用缓存音频: {audio_path}")
            
            if audio_path:
                tts.play_cached(audio_path)
                logger.info(f"🔊 缓存播放完成，总耗时: {time.time() - start_time:.2f}秒")
            else:
                audio_path = tts.speak_sync_with_cache(clean_text)
                if audio_path and last_widget:
                    last_widget._audio_cache = audio_path
                logger.info(f"🔊 合成播放完成，总耗时: {time.time() - start_time:.2f}秒，合成耗时: {time.time() - synth_start:.2f}秒")
            
            self.signal_helper.emit_reset_play_button()
        
        thread = threading.Thread(target=speak_in_thread, daemon=True)
        thread.start()

    def _on_play_video_signal(self, video_path: str, directory: str, video_name: str):
        """Handle play video signal from other threads (runs in main thread)"""
        from loguru import logger
        from pathlib import Path
        import os
        import sys
        import subprocess
        
        logger.info(f"🎬 播放视频: video_path={video_path}, directory={directory}, video_name={video_name}")
        
        try:
            play_path = None
            
            if video_path:
                path = Path(video_path)
                if path.exists():
                    play_path = str(path)
                elif str(video_path).startswith("http"):
                    play_path = video_path
                    
            elif directory:
                dir_path = Path(directory)
                if dir_path.exists() and dir_path.is_dir():
                    video_files = []
                    for ext in ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm', '.m4v']:
                        video_files.extend(dir_path.glob(f"*{ext}"))
                    
                    if video_files:
                        play_path = str(video_files[0])
                        
            elif video_name:
                if video_name.startswith("http"):
                    play_path = video_name
                else:
                    # 尝试作为本地路径
                    path = Path(video_name)
                    if path.exists():
                        play_path = str(path)
            
            if play_path:
                logger.info(f"▶️ 使用系统默认播放器打开: {play_path}")
                
                if sys.platform == "win32":
                    os.startfile(play_path)
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", play_path])
                else:
                    subprocess.Popen(["xdg-open", play_path])
            else:
                logger.error(f"❌ 未找到可播放的视频")
                
        except Exception as e:
            logger.error(f"❌ 播放视频失败: {e}")
            import traceback
            traceback.print_exc()

    async def _show_notification(self, message: str):
        """Show notification - thread-safe version using signals"""
        if threading.current_thread() is threading.main_thread():
            self.append_message("agent", message)
        else:
            self.signal_helper.emit_notification(message)
        return None

    def _on_send(self):
        """Handle send button"""
        from ..tts import get_tts_manager
        tts = get_tts_manager()
        tts.stop()
        self._reset_play_button()
        
        text = self.input_field.text().strip()
        
        if not text and not self.dropped_files:
            return
        
        message_parts = []
        if self.dropped_files:
            file_info = "📎 附件:\n"
            for file_path in self.dropped_files:
                file_name = Path(file_path).name
                file_info += f"  - {file_name}\n"
            message_parts.append(file_info)
        
        if text:
            message_parts.append(text)
        
        full_message = "\n".join(message_parts)
        
        if text:
            self.input_field.add_to_history(text)
        
        self.input_field.clear()
        
        user_metadata = {"files": self.dropped_files.copy()} if self.dropped_files else None
        self.append_message("user", full_message, user_metadata)
        
        files_to_send = self.dropped_files.copy()
        self.dropped_files.clear()
        self._clear_files_preview()
        
        from PyQt6.QtCore import QThread, pyqtSignal
        from PyQt6.QtWidgets import QApplication
        
        chat_history = self._chat_history.copy()
        
        class ResponseThread(QThread):
            response_ready = pyqtSignal(object)
            status_update = pyqtSignal(str)
            web_server_result = pyqtSignal(dict)
            agent_names_ready = pyqtSignal(list)
            chat_history_updated = pyqtSignal(list)
            save_document_ready = pyqtSignal(dict)
            
            def __init__(self, msg, multi_agent, files, chat_history=None):
                super().__init__()
                self.msg = msg
                self.multi_agent = multi_agent
                self.files = files
                self._cancelled = False
                self.chat_history = chat_history or []
            
            def cancel(self):
                self._cancelled = True
            
            def run(self):
                try:
                    from ..channels import IncomingMessage, MessageType
                    
                    self.status_update.emit("💬 AI 正在思考...")
                    
                    context = []
                    for msg in self.chat_history[-10:]:
                        content = msg.get("content", "")
                        if len(content) > 2000:
                            content = content[:2000] + "..."
                        context.append({"role": msg["role"], "content": content})
                    
                    if self.files:
                        files_info = "【附件信息】用户已提供以下附件文件：\n" + "\n".join(f"- {f}" for f in self.files)
                        context.append({"role": "system", "content": files_info})
                    
                    incoming = IncomingMessage(
                        message_id=str(uuid.uuid4()),
                        sender_id="gui_user",
                        sender_name="用户",
                        content=self.msg,
                        message_type=MessageType.TEXT,
                        timestamp=datetime.now(),
                        channel="gui",
                        metadata={
                            "files": self.files,
                            "context": context
                        }
                    )
                    
                    loop = self.multi_agent.loop
                    if loop and not loop.is_closed():
                        future = asyncio.run_coroutine_threadsafe(
                            self.multi_agent.process_message(incoming),
                            loop
                        )
                        
                        import time
                        start_time = time.time()
                        while not future.done():
                            if self._cancelled:
                                future.cancel()
                                self.response_ready.emit("🚫 任务已取消")
                                return
                            elapsed = int(time.time() - start_time)
                            if elapsed > 3:
                                self.status_update.emit(f"⏳ 正在处理中... ({elapsed}秒)")
                            time.sleep(0.5)
                        
                        response = future.result(timeout=180)
                        
                        if response and hasattr(response, 'metadata') and response.metadata:
                            if response.metadata.get('web_server_result'):
                                self.web_server_result.emit(response.metadata['web_server_result'])
                            if response.metadata.get('agent_names'):
                                self.agent_names_ready.emit(response.metadata['agent_names'])
                        
                        if response and hasattr(response, 'content'):
                            result_text = response.content
                        elif response:
                            result_text = str(response)
                        else:
                            result_text = "处理完成"
                        
                        self.chat_history.append({"role": "user", "content": self.msg})
                        self.chat_history.append({"role": "assistant", "content": result_text})
                        if len(self.chat_history) > 40:
                            self.chat_history = self.chat_history[-40:]
                        self.chat_history_updated.emit(self.chat_history.copy())
                        
                        self.response_ready.emit(response)
                    else:
                        self.response_ready.emit("错误: 事件循环未初始化")
                except asyncio.CancelledError:
                    self.response_ready.emit("🚫 任务已取消")
                except Exception as e:
                    import traceback
                    traceback.print_exc()
                    self.response_ready.emit(f"错误: {e}")
        
        def on_response(resp):
            progress_manager.clear_callback()
            
            skip_auto_speak = False
            content = resp
            metadata = None
            
            from loguru import logger
            logger.info(f"🔍 on_response: resp type={type(resp)}, has_metadata={hasattr(resp, 'metadata')}")
            
            if hasattr(resp, 'metadata'):
                metadata = resp.metadata
                if metadata:
                    skip_auto_speak = metadata.get('skip_auto_speak', False)
                    logger.info(f"🔍 on_response: skip_auto_speak={skip_auto_speak}, metadata={metadata}")
                else:
                    logger.info(f"🔍 on_response: metadata is None or empty")
            else:
                logger.info(f"🔍 on_response: resp has no metadata attribute")
            
            if hasattr(resp, 'content'):
                content = resp.content
            
            logger.info(f"🔍 on_response: calling append_message with metadata={metadata}")
            if not skip_auto_speak:
                self.signal_helper.emit_auto_speak(content)
            self.append_message("agent", content, metadata)
        
        def on_chat_history_updated(history):
            self._chat_history = history
        
        def on_status_update(status):
            for i in range(self.messages_list.count() - 1, -1, -1):
                item = self.messages_list.item(i)
                widget = self.messages_list.itemWidget(item)
                if widget and hasattr(widget, '_is_status'):
                    widget.findChild(self.QLabel).setText(status)
                    return
            self.append_message("system", status)
            last_item = self.messages_list.item(self.messages_list.count() - 1)
            last_widget = self.messages_list.itemWidget(last_item)
            if last_widget:
                last_widget._is_status = True
        
        def on_progress(message, progress):
            on_status_update(message)
        
        progress_manager.set_callback(on_progress)
        
        def on_web_server_result(result):
            if result.get('success'):
                self._show_web_server_dialog(
                    result.get('url', ''),
                    result.get('qr_code', '')
                )
        
        def on_agent_names(agent_names):
            from ..config import settings
            base_name = settings.agent.name or "小助手"
            self.chat_title.setText(f"智能助理-{base_name}")
        
        if not hasattr(self, '_response_threads'):
            self._response_threads = []
        
        self._response_threads = [t for t in self._response_threads if t.isRunning()]
        
        new_thread = ResponseThread(text, self.multi_agent, files_to_send, chat_history)
        new_thread.response_ready.connect(on_response)
        new_thread.status_update.connect(on_status_update)
        new_thread.web_server_result.connect(on_web_server_result)
        new_thread.agent_names_ready.connect(on_agent_names)
        new_thread.chat_history_updated.connect(on_chat_history_updated)
        new_thread.finished.connect(lambda: self._cleanup_thread(new_thread))
        new_thread.start()
        self._response_threads.append(new_thread)
    
    def _save_llm_document(self, filename, content):
        """保存 LLM 生成的文档"""
        try:
            from ..config import settings
            import os
            
            download_dir = settings.directories.download_dir
            if not download_dir:
                download_dir = os.path.join(str(Path.home()), "Downloads")
            
            os.makedirs(download_dir, exist_ok=True)
            
            filepath = os.path.join(download_dir, filename)
            
            if filename.endswith('.docx'):
                from docx import Document
                doc = Document()
                
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line:
                        doc.add_paragraph(line)
                
                doc.save(filepath)
                logger.info(f"✅ Word 文档已保存: {filepath}")
                self.append_message("system", f"✅ 文档已保存: {filepath}")
                
            elif filename.endswith('.pdf'):
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfgen import canvas
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.lib.units import cm
                
                c = canvas.Canvas(filepath, pagesize=A4)
                width, height = A4
                
                try:
                    font_path = "C:/Windows/Fonts/simhei.ttf"
                    if os.path.exists(font_path):
                        pdfmetrics.registerFont(TTFont('SimHei', font_path))
                        c.setFont('SimHei', 12)
                except:
                    c.setFont('Helvetica', 12)
                
                y = height - 2*cm
                lines = content.split('\n')
                
                for line in lines:
                    if y < 2*cm:
                        c.showPage()
                        y = height - 2*cm
                        try:
                            c.setFont('SimHei', 12)
                        except:
                            c.setFont('Helvetica', 12)
                    
                    if line.startswith('# '):
                        try:
                            c.setFont('SimHei', 16)
                        except:
                            c.setFont('Helvetica', 16)
                        c.drawString(2*cm, y, line[2:])
                        try:
                            c.setFont('SimHei', 12)
                        except:
                            c.setFont('Helvetica', 12)
                    else:
                        c.drawString(2*cm, y, line)
                    y -= 0.5*cm
                
                c.save()
                logger.info(f"✅ PDF 文档已保存: {filepath}")
                self.append_message("system", f"✅ 文档已保存: {filepath}")
            else:
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(content)
                logger.info(f"✅ 文本文件已保存: {filepath}")
                self.append_message("system", f"✅ 文件已保存: {filepath}")
                
        except Exception as e:
            logger.error(f"保存文档失败: {e}")
            self.append_message("system", f"❌ 保存文档失败: {e}")
    
    def _execute_agent_call(self, alias, task):
        """执行 LLM 触发的智能体调用"""
        try:
            agent_trigger = f"@{alias} {task}"
            logger.info(f"🎯 执行智能体调用: {agent_trigger}")
            
            from ..channels import IncomingMessage, MessageType
            
            incoming = IncomingMessage(
                message_id=str(uuid.uuid4()),
                sender_id="gui_user",
                sender_name="用户",
                content=agent_trigger,
                message_type=MessageType.TEXT,
                timestamp=datetime.now(),
                channel="gui",
                metadata=None
            )
            
            loop = self.multi_agent.loop
            if loop and not loop.is_closed():
                import asyncio
                future = asyncio.run_coroutine_threadsafe(
                    self.multi_agent.process_message(incoming),
                    loop
                )
                response = future.result(timeout=180)
                
                if response and hasattr(response, 'content'):
                    self.append_message("agent", response.content)
                elif response:
                    self.append_message("agent", str(response))
        except Exception as e:
            logger.error(f"执行智能体调用失败: {e}")
            self.append_message("system", f"❌ 智能体调用失败: {e}")
    
    def _cleanup_thread(self, thread):
        """清理完成的线程"""
        if hasattr(self, '_response_threads') and thread in self._response_threads:
            self._response_threads.remove(thread)
        logger.debug(f"Thread cleaned up, {len(self._response_threads) if hasattr(self, '_response_threads') else 0} threads running")
    
    def _on_thread_finished(self):
        """线程完成回调"""
        logger.debug("Response thread finished")

    def _drag_enter_event(self, event):
        """处理拖拽进入事件"""
        from PyQt6.QtCore import Qt
        
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            self.drop_area.setStyleSheet("""
                QWidget {
                    background-color: #e8f5e9;
                    border: 2px dashed #10a37f;
                    border-radius: 12px;
                }
            """)
        else:
            event.ignore()

    def _drop_event(self, event):
        """处理文件/文件夹放下事件"""
        from PyQt6.QtCore import Qt, QUrl
        
        urls = event.mimeData().urls()
        if urls:
            for url in urls:
                file_path = url.toLocalFile()
                if file_path and Path(file_path).exists():
                    self.dropped_files.append(file_path)
                    if Path(file_path).is_dir():
                        logger.info(f"📁 拖入文件夹: {file_path}")
                    else:
                        logger.info(f"📁 拖入文件: {file_path}")
            
            self._update_files_preview()
        
        self.drop_area.setStyleSheet("""
            QWidget {
                background-color: #f7f7f8;
                border: 2px dashed #d9d9e3;
                border-radius: 12px;
            }
        """)
        
        event.acceptProposedAction()

    def _check_agent_trigger(self, text):
        """检查文本是否包含智能体触发器或特殊命令"""
        if not text:
            return False
        
        trigger_patterns = [
            "@音乐智能体", "@音乐", "@music",
            "@视频智能体", "@视频", "@video",
            "@邮件智能体", "@邮件", "@email",
            "@天气智能体", "@天气", "@weather",
            "@文件智能体", "@文件", "@file",
            "@系统智能体", "@系统", "@os",
            "@爬虫智能体", "@爬虫", "@crawler",
            "@下载智能体", "@下载", "@download",
            "@音频解密智能体", "@音频解密", "@audio",
            "@联系人智能体", "@联系人", "@contact",
            "@日历智能体", "@日历", "@calendar",
            "@新闻智能体", "@新闻", "@news",
            "@应用智能体", "@应用", "@app",
            "@开发智能体", "@开发", "@developer",
            "@PDF智能体", "@PDF", "@pdf",
            "@图片智能体", "@图片", "@image",
            "@主动智能体", "@主动", "@proactive",
            "@Web服务智能体", "@Web服务", "@web", "@webserver",
            "@投屏智能体", "@投屏", "@screencast",
            "@股票智能体", "@股票", "@stock",
            "@应用管理智能体", "@应用管理", "@appmgmt",
            "@图片转换智能体", "@图片转换", "@imgconv",
            "@Llm智能体", "@LLM智能体", "@大模型", "@AI", "@ai", "@llm",
            "@旅行智能体", "@旅行", "@travel",
            "@QQ智能体", "@QQ", "@qq",
        ]
        
        text_lower = text.lower().strip()
        
        if text_lower in ["/?", "？", "帮助", "help", "help me"]:
            return True
        
        for pattern in trigger_patterns:
            if pattern.lower() in text_lower:
                return True
        return False

    def _on_browse(self):
        """浏览选择文件或文件夹"""
        from PyQt6.QtWidgets import QFileDialog, QMenu
        from PyQt6.QtCore import QPoint
        
        menu = QMenu(self.window)
        menu.setStyleSheet("""
            QMenu {
                background-color: #ffffff;
                border: 1px solid #d9d9e3;
                border-radius: 8px;
                padding: 4px;
            }
            QMenu::item {
                padding: 8px 16px;
                border-radius: 4px;
            }
            QMenu::item:selected {
                background-color: #e8f5e9;
            }
        """)
        
        file_action = menu.addAction("📄 选择文件")
        folder_action = menu.addAction("📁 选择文件夹")
        
        button_pos = self.browse_button.mapToGlobal(self.browse_button.rect().bottomLeft())
        action = menu.exec(button_pos)
        
        if action == file_action:
            files, _ = QFileDialog.getOpenFileNames(
                self.window,
                "选择文件",
                "",
                "所有文件 (*.*)"
            )
            if files:
                for file_path in files:
                    if file_path not in self.dropped_files:
                        self.dropped_files.append(file_path)
                        logger.info(f"📁 选择文件: {file_path}")
                self._update_files_preview()
        
        elif action == folder_action:
            folder = QFileDialog.getExistingDirectory(
                self.window,
                "选择文件夹",
                ""
            )
            if folder and folder not in self.dropped_files:
                self.dropped_files.append(folder)
                logger.info(f"📁 选择文件夹: {folder}")
                self._update_files_preview()

    def _start_voice_recording(self):
        """开始录音"""
        try:
            from ..utils.voice_input import voice_input_manager
            
            if not voice_input_manager.is_available():
                self.append_message("assistant", voice_input_manager.get_install_hint())
                return
            
            self._voice_recording = True
            self.voice_button.setText("⏹️")
            self.voice_button.setStyleSheet("""
                QPushButton {
                    background-color: #ef4444;
                    color: white;
                    border: 2px solid #dc2626;
                    border-radius: 12px;
                    font-size: 20px;
                }
            """)
            
            def on_status(status: str):
                if status.startswith("error:"):
                    error_msg = status[6:]
                    self._voice_recording = False
                    self.voice_button.setText("🎤")
                    self.voice_button.setStyleSheet("""
                        QPushButton {
                            background-color: #f7f7f8;
                            color: #202123;
                            border: 1px solid #d9d9e3;
                            border-radius: 12px;
                            font-size: 20px;
                        }
                        QPushButton:hover {
                            background-color: #e8f5e9;
                            border: 1px solid #10a37f;
                        }
                    """)
                    self.append_message("agent", error_msg)
                elif status.startswith("warning:"):
                    warning_msg = status[8:]
                    self.append_message("agent", f"⚠️ {warning_msg}")
                elif status == "recording":
                    pass
                elif status == "processing":
                    pass
                elif status == "done":
                    pass
            
            voice_input_manager.set_callbacks(on_status_change=on_status)
            
            if not voice_input_manager.start_recording():
                self._voice_recording = False
                self.voice_button.setText("🎤")
                
        except Exception as e:
            logger.error(f"语音输入错误: {e}")
            self._voice_recording = False
            self.voice_button.setText("🎤")
    
    def _stop_voice_recording(self):
        """停止录音"""
        try:
            from ..utils.voice_input import voice_input_manager
            
            if not self._voice_recording:
                return
            
            self._voice_recording = False
            self.voice_button.setText("🎤")
            self.voice_button.setStyleSheet("""
                QPushButton {
                    background-color: #f7f7f8;
                    color: #202123;
                    border: 1px solid #d9d9e3;
                    border-radius: 12px;
                    font-size: 20px;
                }
                QPushButton:hover {
                    background-color: #e8f5e9;
                    border: 1px solid #10a37f;
                }
            """)
            
            def on_result(text: str):
                pass
            
            voice_input_manager.set_callbacks(on_result=on_result)
            result = voice_input_manager.stop_recording()
            if result:
                current_text = self.input_field.text()
                if current_text:
                    self.input_field.setText(current_text + result)
                else:
                    self.input_field.setText(result)
                    
        except Exception as e:
            logger.error(f"语音输入错误: {e}", exc_info=True)
            self._voice_recording = False
            self.voice_button.setText("🎤")

    def _update_files_preview(self):
        """更新文件预览"""
        while self.files_layout.count():
            item = self.files_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        
        if not self.dropped_files:
            self.files_widget.hide()
            return
        
        self.files_widget.show()
        
        for file_path in self.dropped_files:
            path_obj = Path(file_path)
            file_name = path_obj.name
            is_dir = path_obj.is_dir()
            
            file_widget = self.QWidget()
            file_widget.setStyleSheet("""
                QWidget {
                    background-color: #e3f2fd;
                    border: 1px solid #90caf9;
                    border-radius: 6px;
                }
            """)
            file_layout = self.QHBoxLayout(file_widget)
            file_layout.setContentsMargins(8, 4, 8, 4)
            file_layout.setSpacing(8)
            
            icon_label = self.QLabel("📁" if is_dir else "📄")
            icon_label.setStyleSheet("font-size: 16px;")
            file_layout.addWidget(icon_label)
            
            name_label = self.QLabel(file_name)
            name_label.setStyleSheet("color: #1976d2; font-size: 12px;")
            file_layout.addWidget(name_label)
            
            remove_btn = self.QPushButton("×")
            remove_btn.setFixedSize(20, 20)
            remove_btn.setStyleSheet("""
                QPushButton {
                    background-color: transparent;
                    color: #666;
                    border: none;
                    font-size: 16px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    color: #f44336;
                }
            """)
            remove_btn.clicked.connect(lambda checked, fp=file_path: self._remove_file(fp))
            file_layout.addWidget(remove_btn)
            
            self.files_layout.addWidget(file_widget)
        
        self.files_layout.addStretch()

    def _remove_file(self, file_path: str):
        """移除文件"""
        if file_path in self.dropped_files:
            self.dropped_files.remove(file_path)
            logger.info(f"🗑️ 移除文件: {file_path}")
            self._update_files_preview()

    def _clear_files_preview(self):
        """清空文件预览"""
        self.dropped_files.clear()
        self._update_files_preview()

    def show(self):
        """Show the window"""
        self.window.show()

    def isVisible(self) -> bool:
        """Check if window is visible"""
        return self.window.isVisible()


gui_channel = GUIChannel()
