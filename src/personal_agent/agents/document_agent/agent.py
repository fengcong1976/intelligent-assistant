"""
文档智能体 - 文档处理
支持PDF、Word、Excel文档的读取、生成、转换等功能
"""
import re
from typing import Dict, Any, Optional, List, Union
from pathlib import Path
from loguru import logger

from ..base import BaseAgent, Task


class DocumentAgent(BaseAgent):
    """文档智能体 - 处理PDF、Word、Excel文档的读取、生成和转换"""
    
    KEYWORD_MAPPINGS = {
        "读取PDF": ("pdf_read", {}),
        "打开PDF": ("pdf_read", {}),
        "PDF摘要": ("pdf_summarize", {}),
        "PDF总结": ("pdf_summarize", {}),
        "提取PDF文本": ("pdf_extract_text", {}),
        "PDF转Word": ("pdf_to_word", {}),
        "PDF转图片": ("pdf_to_image", {}),
        "Word转PDF": ("word_to_pdf", {}),
        "合并PDF": ("pdf_merge", {}),
        "拆分PDF": ("pdf_split", {}),
        "生成PDF": ("pdf_generate", {}),
        "生成文档": ("doc_generate", {}),
        "生成Word": ("doc_generate", {}),
        "生成DOC": ("doc_generate", {}),
        "生成doc": ("doc_generate", {}),
        "存成doc": ("doc_generate", {}),
        "保存为doc": ("doc_generate", {}),
        "保存为Word": ("doc_generate", {}),
        "保存为word": ("doc_generate", {}),
        "生成Excel": ("excel_generate", {}),
        "生成excel": ("excel_generate", {}),
        "生成表格": ("excel_generate", {}),
        "存成Excel": ("excel_generate", {}),
        "保存为Excel": ("excel_generate", {}),
        "创建表格": ("excel_generate", {}),
    }
    
    def __init__(self):
        super().__init__(
            name="document_agent",
            description="文档智能体 - 支持PDF、Word、Excel文档的浏览、解析、生成和格式转换"
        )
        
        self.register_capability(
            capability="save_document",
            description="将内容保存为文档文件。支持 Word (.docx)、Excel (.xlsx) 和 PDF 格式。根据文件扩展名自动选择格式。注意：content 可以为空，表示需要从前序工具获取内容。",
            parameters={
                "type": "object",
                "properties": {
                    "content": {
                        "type": "string",
                        "description": "文档内容。可以为空，系统会自动从前序工具获取内容"
                    },
                    "filename": {
                        "type": "string",
                        "description": "文件名，如'报告.docx'、'数据.xlsx'或'文档.pdf'"
                    }
                },
                "required": ["filename"]
            },
            category="document"
        )
        
        self.register_capability("pdf_read", "读取PDF")
        self.register_capability("pdf_extract_text", "提取PDF文本")
        self.register_capability("pdf_summarize", "PDF摘要")
        self.register_capability("pdf_generate", "生成PDF")
        self.register_capability("doc_generate", "生成Word文档")
        self.register_capability("excel_generate", "生成Excel")
        self.register_capability("pdf_merge", "合并PDF")
        self.register_capability("pdf_split", "拆分PDF")
        self.register_capability("pdf_to_word", "PDF转Word")
        self.register_capability("pdf_to_image", "PDF转图片")
        self.register_capability("word_to_pdf", "Word转PDF")
        
        self.register_file_formats(
            open_formats=[".pdf", ".docx", ".doc", ".xlsx", ".xls"],
            edit_formats=[".pdf", ".docx", ".xlsx", ".png", ".jpg"]
        )
        
        self._llm_gateway = None
        
        logger.info("📄 文档智能体已初始化")

    def _get_llm_gateway(self):
        """获取 LLM 网关"""
        if self._llm_gateway is None:
            from ...llm import LLMGateway
            from ...config import settings
            self._llm_gateway = LLMGateway(settings.llm)
        return self._llm_gateway

    def _get_documents_dir(self) -> Path:
        """获取文档保存目录"""
        try:
            from ...config import settings
            return settings.directory.get_documents_dir()
        except Exception as e:
            logger.warning(f"获取文档目录失败: {e}")
            return Path.cwd() / "output" / "documents"

    async def _call_llm(self, prompt: str) -> str:
        """调用 LLM"""
        llm = self._get_llm_gateway()
        messages = [{"role": "user", "content": prompt}]
        response = await llm.chat(messages)
        return response.content

    async def execute_task(self, task: Task) -> Any:
        task_type = task.type
        params = task.params
        logger.info(f"📄 文档智能体执行任务: {task_type}")
        
        if task_type == "general":
            return await self._handle_general(params)
        elif task_type == "save_document":
            return await self._handle_save_document(params)
        elif task_type == "pdf_read":
            return await self._handle_read(params)
        elif task_type == "pdf_extract_text":
            return await self._handle_extract_text(params)
        elif task_type == "pdf_summarize":
            return await self._handle_summarize(params)
        elif task_type == "pdf_generate":
            return await self._handle_generate(params)
        elif task_type == "pdf_merge":
            return await self._handle_merge(params)
        elif task_type == "pdf_split":
            return await self._handle_split(params)
        elif task_type == "pdf_to_word":
            return await self._handle_pdf_to_word(params)
        elif task_type == "pdf_to_image":
            return await self._handle_pdf_to_image(params)
        elif task_type == "word_to_pdf":
            return await self._handle_word_to_pdf(params)
        elif task_type == "txt_to_pdf":
            return await self._handle_txt_to_pdf(params)
        elif task_type == "image_to_pdf":
            return await self._handle_image_to_pdf(params)
        elif task_type == "excel_to_pdf":
            return await self._handle_excel_to_pdf(params)
        elif task_type == "convert":
            return await self._handle_convert(params)
        elif task_type == "doc_generate":
            return await self._handle_doc_generate(params)
        elif task_type == "excel_generate":
            return await self._handle_excel_generate(params)
        elif task_type == "agent_help":
            return self._get_help_info()
        else:
            return f"❌ 不支持的任务类型: {task_type}"

    async def _handle_general(self, params: Dict) -> str:
        """处理自然语言请求"""
        original_text = params.get("original_text", "") or params.get("text", "")
        file_path = params.get("file_path", "")
        
        if "保存" in original_text and any(fmt in original_text.lower() for fmt in ["pdf", "word", "doc", "excel", "xlsx", "txt"]):
            content = params.get("content", "")
            filename = params.get("filename", "")
            
            if not content and "写" in original_text:
                topic_match = re.search(r"关于(.+?)的?文章", original_text)
                topic = topic_match.group(1) if topic_match else "主题"
                
                prompt = f"请写一篇关于{topic}的文章，要求内容丰富、结构清晰，大约500-800字。"
                content = await self._call_llm(prompt)
            
            if not filename:
                if "pdf" in original_text.lower():
                    filename = "文档.pdf"
                elif "word" in original_text.lower() or "doc" in original_text.lower():
                    filename = "文档.docx"
                elif "excel" in original_text.lower() or "xlsx" in original_text.lower():
                    filename = "文档.xlsx"
                else:
                    filename = "文档.txt"
            
            save_params = {"filename": filename, "content": content}
            return await self._handle_save_document(save_params)
        
        if not file_path:
            return "❌ 请提供要转换的文件"
        
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if "转pdf" in original_text or "转成pdf" in original_text or "转换为pdf" in original_text:
            if suffix in [".docx", ".doc"]:
                return await self._handle_word_to_pdf({"path": file_path})
            elif suffix == ".txt":
                return await self._handle_txt_to_pdf({"path": file_path})
            elif suffix in [".png", ".jpg", ".jpeg", ".bmp", ".gif"]:
                return await self._handle_image_to_pdf({"path": file_path})
            elif suffix in [".xlsx", ".xls"]:
                return await self._handle_excel_to_pdf({"path": file_path})
            elif task_type == "agent_help":
                return self._get_help_info()
            else:
                return f"❌ 不支持的文件格式: {suffix}"
        elif "转word" in original_text or "转成word" in original_text or "转换为word" in original_text:
            return await self._handle_pdf_to_word({"path": file_path})
        elif "转图片" in original_text or "转成图片" in original_text:
            return await self._handle_pdf_to_image({"path": file_path})
        elif "读取" in original_text or "打开" in original_text:
            return await self._handle_read({"path": file_path or original_text})
        elif "提取" in original_text or "文本" in original_text:
            return await self._handle_extract_text({"path": file_path or original_text})
        elif "摘要" in original_text or "总结" in original_text:
            return await self._handle_summarize({"path": file_path or original_text})
        elif "生成" in original_text or "创建" in original_text:
            return await self._handle_generate({"content": original_text})
        elif "合并" in original_text:
            return await self._handle_merge({"paths": []})
        elif "拆分" in original_text or "分割" in original_text:
            return await self._handle_split({"path": file_path or original_text})
        else:
            return """📄 文档智能体可以帮助你：

**PDF操作**:
1. **读取PDF**: 读取PDF文件内容
2. **提取文本**: 从PDF中提取纯文本
3. **生成摘要**: 自动生成PDF内容摘要
4. **生成PDF**: 从文本或内容创建PDF文件
5. **合并PDF**: 合并多个PDF文件
6. **拆分PDF**: 将PDF拆分为多个文件

**Word操作**:
7. **生成Word**: 创建Word文档（.docx）
8. **Word转PDF**: 将Word文档转换为PDF

**Excel操作**:
9. **生成Excel**: 创建Excel表格（.xlsx）
10. **Excel转PDF**: 将Excel表格转换为PDF

**格式转换**:
11. **PDF转Word**: 将PDF转换为Word文档
12. **PDF转图片**: 将PDF页面转换为图片
13. **TXT转PDF**: 将文本文件转换为PDF
14. **图片转PDF**: 将图片转换为PDF

请告诉我你需要什么操作？"""

    def _check_pypdf(self):
        """检查 pypdf 是否可用"""
        try:
            import pypdf
            return True
        except ImportError:
            return False

    def _check_reportlab(self):
        """检查 reportlab 是否可用"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            return True
        except ImportError:
            return False

    def _check_python_docx(self) -> bool:
        """检查 python-docx 是否可用"""
        try:
            from docx import Document
            return True
        except ImportError:
            return False

    def _check_openpyxl(self) -> bool:
        """检查 openpyxl 是否可用"""
        try:
            from openpyxl import Workbook
            return True
        except ImportError:
            return False

    async def _handle_read(self, params: Dict) -> str:
        """读取PDF文件"""
        path = params.get("path", "") or params.get("file_path", "")
        
        if not path:
            return "❌ 请提供PDF文件路径"
        
        path = Path(path)
        if not path.exists():
            return f"❌ 文件不存在: {path}"
        
        if not self._check_pypdf():
            return "❌ 需要安装 pypdf 库: pip install pypdf"
        
        try:
            import pypdf
            
            reader = pypdf.PdfReader(str(path))
            num_pages = len(reader.pages)
            
            info = reader.metadata
            info_text = ""
            if info:
                if info.title:
                    info_text += f"标题: {info.title}\n"
                if info.author:
                    info_text += f"作者: {info.author}\n"
                if info.subject:
                    info_text += f"主题: {info.subject}\n"
            
            text_preview = ""
            if num_pages > 0:
                first_page_text = reader.pages[0].extract_text()
                text_preview = first_page_text[:500] + "..." if len(first_page_text) > 500 else first_page_text
            
            return f"""📄 PDF文件信息

📁 文件: {path.name}
📊 页数: {num_pages}
{info_text}

📝 内容预览:
{text_preview}"""
        except Exception as e:
            return f"❌ 读取PDF失败: {e}"

    async def _handle_extract_text(self, params: Dict) -> str:
        """提取PDF文本"""
        path = params.get("path", "") or params.get("file_path", "")
        output_path = params.get("output", "")
        
        if not path:
            return "❌ 请提供PDF文件路径"
        
        path = Path(path)
        if not path.exists():
            return f"❌ 文件不存在: {path}"
        
        if not self._check_pypdf():
            return "❌ 需要安装 pypdf 库: pip install pypdf"
        
        try:
            import pypdf
            
            reader = pypdf.PdfReader(str(path))
            all_text = []
            
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    all_text.append(f"=== 第 {i} 页 ===\n{text}")
            
            full_text = "\n\n".join(all_text)
            
            if output_path:
                output_path = Path(output_path)
                output_path.write_text(full_text, encoding='utf-8')
                return f"✅ 文本已提取并保存到: {output_path}"
            
            preview = full_text[:1000] + "..." if len(full_text) > 1000 else full_text
            return f"""✅ 提取完成

📊 总页数: {len(reader.pages)}
📝 文本长度: {len(full_text)} 字符

预览:
{preview}"""
        except Exception as e:
            return f"❌ 提取文本失败: {e}"

    async def _handle_summarize(self, params: Dict) -> str:
        """生成PDF摘要"""
        path = params.get("path", "") or params.get("file_path", "")
        
        if not path:
            return "❌ 请提供PDF文件路径"
        
        path = Path(path)
        if not path.exists():
            return f"❌ 文件不存在: {path}"
        
        if not self._check_pypdf():
            return "❌ 需要安装 pypdf 库: pip install pypdf"
        
        try:
            import pypdf
            
            reader = pypdf.PdfReader(str(path))
            all_text = []
            
            for page in reader.pages[:10]:
                text = page.extract_text()
                if text.strip():
                    all_text.append(text)
            
            content = "\n\n".join(all_text)
            
            if len(content) > 8000:
                content = content[:8000] + "..."
            
            prompt = f"""请为以下PDF内容生成摘要：

{content}

请提供：
1. 主要内容概述（100字以内）
2. 关键要点（3-5个）
3. 适合的标签"""
            
            summary = await self._call_llm(prompt)
            
            return f"""📄 PDF摘要

📁 文件: {path.name}
📊 页数: {len(reader.pages)}

{summary}"""
        except Exception as e:
            return f"❌ 生成摘要失败: {e}"

    async def _handle_generate(self, params: Dict) -> str:
        """生成PDF文件"""
        content = params.get("content", "")
        title = params.get("title", "文档")
        output_path = params.get("output", "")
        
        if not content:
            return "❌ 请提供PDF内容"
        
        if len(content) < 100:
            logger.info(f"📝 内容较短，调用LLM生成详细内容: {title}")
            generated_content = await self._generate_content_with_llm(title, content)
            if generated_content:
                content = generated_content
        
        output_dir = self._get_documents_dir()
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if not output_path:
            output_path = f"{title}.pdf"
        
        output_path = Path(output_path)
        if not output_path.is_absolute():
            output_path = output_dir / output_path.name
        
        return await self._generate_with_reportlab(title, content, output_path)

    def _register_chinese_font(self) -> str:
        """注册中文字体"""
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        font_paths = [
            ("C:/Windows/Fonts/simhei.ttf", "SimHei"),
            ("C:/Windows/Fonts/msyh.ttc", "MicrosoftYaHei"),
            ("C:/Windows/Fonts/simsun.ttc", "SimSun"),
            ("C:/Windows/Fonts/simkai.ttf", "KaiTi"),
        ]
        
        for font_path, font_name in font_paths:
            if Path(font_path).exists():
                try:
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    logger.info(f"✅ 已注册中文字体: {font_name}")
                    return font_name
                except Exception as e:
                    logger.warning(f"注册字体失败 {font_name}: {e}")
                    continue
        
        logger.warning("未找到中文字体，使用默认字体")
        return "Helvetica"

    async def _generate_with_reportlab(self, title: str, content: str, output_path: Path) -> str:
        """使用reportlab生成PDF"""
        if not self._check_reportlab():
            return "❌ 需要安装 reportlab 库: pip install reportlab"
        
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
            import re
            
            font_name = self._register_chinese_font()
            
            styles = getSampleStyleSheet()
            
            title_style = ParagraphStyle(
                'ChineseTitle',
                parent=styles['Title'],
                fontName=font_name,
                fontSize=22,
                alignment=TA_CENTER,
                spaceAfter=20,
                spaceBefore=10,
            )
            
            h2_style = ParagraphStyle(
                'ChineseH2',
                parent=styles['Heading2'],
                fontName=font_name,
                fontSize=16,
                spaceBefore=15,
                spaceAfter=8,
            )
            
            h3_style = ParagraphStyle(
                'ChineseH3',
                parent=styles['Heading3'],
                fontName=font_name,
                fontSize=14,
                spaceBefore=12,
                spaceAfter=6,
            )
            
            body_style = ParagraphStyle(
                'ChineseBody',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
                leading=18,
                alignment=TA_JUSTIFY,
                spaceBefore=4,
                spaceAfter=4,
            )
            
            list_style = ParagraphStyle(
                'ChineseList',
                parent=body_style,
                leftIndent=20,
            )
            
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4,
                leftMargin=2*cm,
                rightMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm,
            )
            
            story = []
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.5*cm))
            
            lines = content.split('\n')
            for line in lines:
                stripped = line.strip()
                
                if not stripped:
                    story.append(Spacer(1, 0.3*cm))
                    continue
                
                text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', stripped)
                text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
                
                if stripped.startswith('# ') and not stripped.startswith('## '):
                    story.append(Paragraph(stripped[2:], title_style))
                elif stripped.startswith('## ') and not stripped.startswith('### '):
                    story.append(Paragraph(stripped[3:], h2_style))
                elif stripped.startswith('### '):
                    story.append(Paragraph(stripped[4:], h3_style))
                elif stripped.startswith('- ') or stripped.startswith('* '):
                    story.append(Paragraph('• ' + text[2:], list_style))
                elif re.match(r'^\d+\.\s', stripped):
                    story.append(Paragraph(text, list_style))
                else:
                    story.append(Paragraph(text, body_style))
            
            doc.build(story)
            
            logger.info(f"✅ PDF已生成: {output_path}")
            return f"✅ PDF已生成\n\n📁 保存位置: {output_path}\n\n💡 可以在文件管理器中查看"
        except Exception as e:
            logger.error(f"PDF生成失败: {e}")
            import traceback
            traceback.print_exc()
            return f"❌ 生成PDF失败: {e}"

    async def _generate_content_with_llm(self, title: str, hint: str) -> str:
        """使用LLM生成详细内容"""
        llm = self._get_llm_gateway()
        if not llm:
            return ""
        
        prompt = f"""你是一位专业的文档撰写专家。请为以下主题生成一份详细、专业、内容丰富的文档。

主题：{title}
提示：{hint}

请生成完整的内容，要求：
1. 标题要吸引人，能准确概括主题
2. 开篇要有引人入胜的引言，概述主题的重要性和背景
3. 正文内容要分章节、分层次展开，每个章节要有：
   - 清晰的小标题
   - 详细的说明和论述
   - 具体的案例、数据或事实支撑
   - 每个要点要展开说明，不要只是简单列举
4. 内容要有深度，避免泛泛而谈，要提供有价值的见解
5. 语言要专业但通俗易懂，适合大众阅读
6. 结尾要有总结和展望，给读者留下深刻印象
7. 全文总字数建议在 1500-2500 字

请直接输出文档内容，使用以下格式：
# 主标题

## 一、章节标题
正文内容...

### 1.1 小节标题
详细内容...

## 二、章节标题
...

## 总结
总结内容..."""

        try:
            messages = [{"role": "user", "content": prompt}]
            response = await llm.chat(messages)
            content = response.content.strip()
            logger.info(f"✅ LLM生成内容完成，共 {len(content)} 字符")
            return content
        except Exception as e:
            logger.error(f"LLM生成内容失败: {e}")
            return ""

    async def _handle_doc_generate(self, params: Dict) -> str:
        """生成Word文档（.docx格式）"""
        content = params.get("content", "")
        title = params.get("title", "文档")
        output_path = params.get("output", "")
        
        if not content:
            logger.info(f"📝 内容为空，调用LLM生成详细内容: {title}")
            content = await self._generate_content_with_llm(title, "")
            if not content:
                return "❌ 无法生成文档内容"
        
        if len(content) < 100:
            logger.info(f"📝 内容较短，调用LLM生成详细内容: {title}")
            generated_content = await self._generate_content_with_llm(title, content)
            if generated_content:
                content = generated_content
        
        output_dir = self._get_documents_dir()
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if not output_path:
            output_path = f"{title}.docx"
        
        output_path = Path(output_path)
        if not output_path.is_absolute():
            output_path = output_dir / output_path.name
        
        if output_path.suffix.lower() == ".doc":
            output_path = output_path.with_suffix(".docx")
        
        return await self._generate_docx(title, content, output_path)

    async def _handle_save_document(self, params: Dict) -> Union[str, Dict[str, Any]]:
        """保存文档 - 根据文件扩展名自动选择格式"""
        content = params.get("content", "")
        filename = params.get("filename", "文档.docx")
        
        output_dir = self._get_documents_dir()
        output_dir.mkdir(parents=True, exist_ok=True)
        
        output_path = Path(filename)
        if not output_path.is_absolute():
            output_path = output_dir / output_path.name
        
        suffix = output_path.suffix.lower()
        title = output_path.stem
        
        excel_keywords = ["报表", "表格", "数据表", "信息表", "通讯录", "名单", "清单", "列表", "财务", "销售", "库存", "员工", "产品", "价格"]
        is_excel_like = any(kw in title for kw in excel_keywords)
        
        if is_excel_like:
            suffix = ".xlsx"
            output_path = output_path.with_suffix(suffix)
        elif not suffix:
            suffix = ".docx"
            output_path = output_path.with_suffix(suffix)
        
        if not content:
            if suffix in [".xlsx", ".xls"]:
                content = title
            else:
                content = await self._generate_content_with_llm(title, "")
        
        result = None
        if suffix in [".xlsx", ".xls"]:
            result = await self._generate_excel(content, output_path)
        elif suffix == ".pdf":
            result = await self._generate_pdf_from_content(content, output_path)
        else:
            if suffix == ".doc":
                output_path = output_path.with_suffix(".docx")
            result = await self._generate_docx(title, content, output_path)
        
        if isinstance(result, str) and "✅" in result:
            return {
                "message": result,
                "file_path": str(output_path),
                "filename": output_path.name,
                "format": suffix[1:] if suffix else "docx"
            }
        
        return result

    async def _generate_excel(self, content: str, output_path: Path) -> str:
        """生成Excel文件"""
        title = output_path.stem if output_path else "数据表格"
        
        parsed = self._parse_excel_content(content) if content else None
        headers = []
        data = []
        
        if parsed and parsed.get("data") and len(parsed.get("data", [])) > 1:
            headers = parsed.get("headers", [])
            data = parsed.get("data", [])
        else:
            logger.info(f"📊 数据为空或不足，调用LLM生成表格数据: {title}")
            generated = await self._generate_excel_data_with_llm(title, content)
            if generated and generated.get("data"):
                headers = generated.get("headers", [])
                data = generated.get("data", [])
        
        return await self._generate_xlsx(title, headers, data, output_path)

    async def _generate_pdf_from_content(self, content: str, output_path: Path) -> str:
        """从内容生成PDF"""
        return await self._handle_generate({"content": content, "output": str(output_path)})

    async def _generate_docx(self, title: str, content: str, output_path: Path) -> str:
        """使用 python-docx 生成 Word 文档"""
        if not self._check_python_docx():
            return "❌ 需要安装 python-docx 库: pip install python-docx"
        
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re
            
            doc = Document()
            
            style = doc.styles['Normal']
            font = style.font
            font.name = 'SimSun'
            font.size = Pt(12)
            
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            lines = content.split('\n')
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    continue
                
                if stripped.startswith('# ') and not stripped.startswith('## '):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith('## ') and not stripped.startswith('### '):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith('### '):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith('- ') or stripped.startswith('* '):
                    doc.add_paragraph(stripped[2:], style='List Bullet')
                elif re.match(r'^\d+\.\s', stripped):
                    doc.add_paragraph(stripped, style='List Number')
                elif stripped.startswith('**') and stripped.endswith('**'):
                    para = doc.add_paragraph()
                    run = para.add_run(stripped[2:-2])
                    run.bold = True
                else:
                    doc.add_paragraph(stripped)
            
            doc.save(str(output_path))
            
            logger.info(f"✅ Word文档已生成: {output_path}")
            return f"✅ Word文档已生成\n\n📁 保存位置: {output_path}\n\n💡 可以用 Word 或 WPS 打开查看"
        except Exception as e:
            logger.error(f"Word文档生成失败: {e}")
            import traceback
            traceback.print_exc()
            return f"❌ 生成Word文档失败: {e}"

    async def _handle_excel_generate(self, params: Dict) -> str:
        """生成Excel表格（.xlsx格式）"""
        content = params.get("content", "") or params.get("data", "")
        title = params.get("title", "")
        headers = params.get("headers", [])
        data = params.get("data", [])
        output_path = params.get("output", "")
        
        if isinstance(data, str):
            content = data
            data = []
        
        if not title:
            if "通讯录" in content or "联系人" in content:
                title = "通讯录"
            elif "邮箱" in content:
                title = "邮箱列表"
            else:
                title = "数据表格"
        
        output_dir = self._get_documents_dir()
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if not output_path:
            output_path = f"{title}.xlsx"
        
        output_path = Path(output_path)
        if not output_path.is_absolute():
            output_path = output_dir / output_path.name
        
        if output_path.suffix.lower() == ".xls":
            output_path = output_path.with_suffix(".xlsx")
        
        if content and (not data or isinstance(data, str)):
            if isinstance(data, str):
                content = data
            parsed = self._parse_excel_content(content)
            if parsed and parsed.get("data") and len(parsed.get("data", [])) > 1:
                headers = parsed.get("headers", headers)
                data = parsed.get("data", data)
        
        if not data or (isinstance(data, list) and len(data) == 0) or (isinstance(data, list) and len(data) <= 1):
            logger.info(f"📊 数据为空或不足，调用LLM生成表格数据: {title}")
            generated = await self._generate_excel_data_with_llm(title, content)
            if generated and generated.get("data"):
                headers = generated.get("headers", headers)
                data = generated.get("data", data)
        
        return await self._generate_xlsx(title, headers, data, output_path)
    
    async def _generate_excel_data_with_llm(self, title: str, hint: str) -> Dict:
        """使用LLM生成Excel表格数据"""
        llm = self._get_llm_gateway()
        if not llm:
            return {}
        
        prompt = f"""你是一位数据整理专家。请根据以下主题生成一份详细的Excel表格数据。

主题：{title}
提示：{hint}

请生成表格数据，要求：
1. 表头要清晰、专业，能准确描述每列数据
2. 数据要真实、有代表性，至少包含 10-15 行数据
3. 每行数据要完整，不要有空值
4. 数据要有实际意义，能够体现主题特点

请以JSON格式返回，格式如下：
{{
    "headers": ["列名1", "列名2", "列名3", ...],
    "data": [
        ["数据1", "数据2", "数据3", ...],
        ["数据1", "数据2", "数据3", ...],
        ...
    ]
}}

请直接返回JSON，不要包含其他文字。"""

        try:
            import json
            messages = [{"role": "user", "content": prompt}]
            response = await llm.chat(messages)
            content = response.content.strip()
            
            if content.startswith("```"):
                lines = content.split("\n")
                content = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])
            
            result = json.loads(content)
            logger.info(f"✅ LLM生成Excel数据完成，{len(result.get('data', []))} 行")
            return result
        except Exception as e:
            logger.error(f"LLM生成Excel数据失败: {e}")
            return {}

    def _parse_excel_content(self, content: str) -> Dict:
        """解析内容为表格数据"""
        result = {"headers": [], "data": []}
        
        content = content.strip()
        if not content:
            return result
        
        try:
            import json
            data = json.loads(content)
            if isinstance(data, list) and data:
                if isinstance(data[0], dict):
                    result["headers"] = list(data[0].keys())
                    result["data"] = [[item.get(h, "") for h in result["headers"]] for item in data]
                    return result
        except (json.JSONDecodeError, TypeError):
            pass
        
        lines = [l.strip() for l in content.split('\n') if l.strip()]
        
        if not lines:
            return result
        
        import re
        contact_rows = []
        
        for line in lines:
            if '📧' in line or '📞' in line:
                name_match = re.search(r'•\s*(.+?)(?:\s*[📞📧👥📋]|$)', line)
                email_match = re.search(r'📧\s*(\S+)', line)
                phone_match = re.search(r'📞\s*(\S+)', line)
                relation_match = re.search(r'👥\s*(\S+)', line)
                
                if name_match:
                    name = name_match.group(1).strip()
                    email = email_match.group(1).strip() if email_match else ""
                    phone = phone_match.group(1).strip() if phone_match else ""
                    relation = relation_match.group(1).strip() if relation_match else ""
                    contact_rows.append([name, email, phone, relation])
        
        if contact_rows:
            result["headers"] = ["姓名", "邮箱", "电话", "关系"]
            result["data"] = contact_rows
            return result
        
        for sep in ['|', '\t', ',']:
            if sep in lines[0]:
                headers = [h.strip() for h in lines[0].split(sep) if h.strip()]
                if len(headers) >= 2:
                    result["headers"] = headers
                    for line in lines[1:]:
                        row = [c.strip() for c in line.split(sep)]
                        if row and any(row):
                            while len(row) < len(headers):
                                row.append("")
                            result["data"].append(row[:len(headers)])
                    return result
        
        result["data"] = [[line] for line in lines]
        result["headers"] = ["内容"]
        return result

    async def _generate_xlsx(self, title: str, headers: List[str], data: List[List], output_path: Path) -> str:
        """使用 openpyxl 生成 Excel 文件"""
        if not self._check_openpyxl():
            return "❌ 需要安装 openpyxl 库: pip install openpyxl"
        
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
            from openpyxl.utils import get_column_letter
            
            wb = Workbook()
            ws = wb.active
            ws.title = title[:31] if len(title) > 31 else title
            
            header_font = Font(bold=True, size=12)
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font_white = Font(bold=True, size=12, color="FFFFFF")
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            center_align = Alignment(horizontal='center', vertical='center')
            
            if headers:
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = header_font_white
                    cell.fill = header_fill
                    cell.border = thin_border
                    cell.alignment = center_align
            
            for row_idx, row_data in enumerate(data, start=2 if headers else 1):
                for col_idx, value in enumerate(row_data, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.border = thin_border
                    cell.alignment = Alignment(vertical='center')
            
            for col in range(1, max(len(headers), max(len(row) for row in data) if data else 1) + 1):
                max_length = 0
                column_letter = get_column_letter(col)
                for row in range(1, ws.max_row + 1):
                    cell_value = ws.cell(row=row, column=col).value
                    if cell_value:
                        max_length = max(max_length, len(str(cell_value)))
                ws.column_dimensions[column_letter].width = min(max_length + 2, 50)
            
            wb.save(str(output_path))
            
            row_count = len(data)
            col_count = len(headers) if headers else (max(len(row) for row in data) if data else 0)
            
            logger.info(f"✅ Excel已生成: {output_path}")
            return f"✅ 已生成Excel文档\n文件路径: {output_path}"
        except Exception as e:
            logger.error(f"Excel生成失败: {e}")
            import traceback
            traceback.print_exc()
            return f"❌ 生成Excel失败: {e}"

    async def _handle_merge(self, params: Dict) -> str:
        """合并PDF文件"""
        paths = params.get("paths", [])
        output_path = params.get("output", "merged.pdf")
        
        if not paths:
            return "❌ 请提供要合并的PDF文件列表"
        
        if not self._check_pypdf():
            return "❌ 需要安装 pypdf 库: pip install pypdf"
        
        try:
            import pypdf
            
            merger = pypdf.PdfMerger()
            
            for path in paths:
                path = Path(path)
                if path.exists():
                    merger.append(str(path))
                else:
                    return f"❌ 文件不存在: {path}"
            
            merger.write(output_path)
            merger.close()
            
            return f"✅ 已合并 {len(paths)} 个PDF文件到: {output_path}"
        except Exception as e:
            return f"❌ 合并PDF失败: {e}"

    async def _handle_split(self, params: Dict) -> str:
        """拆分PDF文件"""
        path = params.get("path", "") or params.get("file_path", "")
        pages = params.get("pages", "")
        output_dir = params.get("output_dir", ".")
        
        if not path:
            return "❌ 请提供PDF文件路径"
        
        path = Path(path)
        if not path.exists():
            return f"❌ 文件不存在: {path}"
        
        if not self._check_pypdf():
            return "❌ 需要安装 pypdf 库: pip install pypdf"
        
        try:
            import pypdf
            
            reader = pypdf.PdfReader(str(path))
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            
            if pages:
                page_list = [int(p.strip()) - 1 for p in pages.split(',') if p.strip().isdigit()]
            else:
                page_list = range(len(reader.pages))
            
            created_files = []
            for i in page_list:
                if 0 <= i < len(reader.pages):
                    writer = pypdf.PdfWriter()
                    writer.add_page(reader.pages[i])
                    
                    output_file = output_dir / f"{path.stem}_page_{i+1}.pdf"
                    with open(output_file, 'wb') as f:
                        writer.write(f)
                    created_files.append(str(output_file))
            
            return f"""✅ PDF已拆分

📁 原文件: {path.name}
📄 创建文件: {len(created_files)} 个

{chr(10).join(created_files[:10])}"""
        except Exception as e:
            return f"❌ 拆分PDF失败: {e}"

    def _check_pdf2docx(self):
        """检查 pdf2docx 是否可用"""
        try:
            from pdf2docx import Converter
            return True
        except ImportError:
            return False

    def _check_pdf2image(self):
        """检查 pdf2image 是否可用"""
        try:
            from pdf2image import convert_from_path
            return True
        except ImportError:
            return False

    async def _handle_convert(self, params: Dict) -> str:
        """处理通用转换请求"""
        target_format = params.get("target_format", "word").lower()
        path = params.get("path", "") or params.get("file_path", "")
        
        if target_format in ["word", "docx"]:
            return await self._handle_pdf_to_word(params)
        elif target_format in ["image", "png", "jpg", "jpeg"]:
            return await self._handle_pdf_to_image(params)
        elif task_type == "agent_help":
            return self._get_help_info()
        else:
            return f"❌ 不支持的目标格式: {target_format}"

    async def _handle_pdf_to_word(self, params: Dict) -> str:
        """将PDF转换为Word文档"""
        path = params.get("path", "") or params.get("file_path", "")
        output_dir = params.get("output_dir", "")
        
        if not path:
            return "❌ 请提供PDF文件路径"
        
        path = Path(path)
        if not path.exists():
            return f"❌ 文件不存在: {path}"
        
        if path.suffix.lower() != '.pdf':
            return f"❌ 不是PDF文件: {path}"
        
        if not self._check_pdf2docx():
            return "❌ 需要安装 pdf2docx 库: pip install pdf2docx"
        
        try:
            from pdf2docx import Converter
            import asyncio
            
            output_dir = Path(output_dir) if output_dir else path.parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            output_path = output_dir / f"{path.stem}.docx"
            
            if output_path.exists():
                return f"⚠️ 目标文件已存在: {output_path}"
            
            logger.info(f"开始转换PDF到Word: {path} → {output_path}")
            
            cv = Converter(str(path))
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, lambda: cv.convert(str(output_path), start=0, end=None))
            cv.close()
            
            logger.info(f"✅ PDF转Word成功: {output_path}")
            return f"✅ 已转换为Word文档\n\n📁 保存位置: {output_path}"
        
        except Exception as e:
            if "password" in str(e).lower():
                return "🔒 PDF受密码保护，请提供密码"
            logger.error(f"PDF转Word失败: {e}")
            return f"❌ 转换失败: {str(e)}"

    async def _handle_pdf_to_image(self, params: Dict) -> str:
        """将PDF转换为图片"""
        path = params.get("path", "") or params.get("file_path", "")
        output_dir = params.get("output_dir", "")
        pages = params.get("pages", "")
        dpi = params.get("dpi", 200)
        
        if not path:
            return "❌ 请提供PDF文件路径"
        
        path = Path(path)
        if not path.exists():
            return f"❌ 文件不存在: {path}"
        
        if path.suffix.lower() != '.pdf':
            return f"❌ 不是PDF文件: {path}"
        
        if not self._check_pdf2image():
            return "❌ 需要安装 pdf2image 库: pip install pdf2image\n\n💡 还需要安装 poppler"
        
        try:
            from pdf2image import convert_from_path
            import asyncio
            
            output_dir = Path(output_dir) if output_dir else path.parent / f"{path.stem}_images"
            output_dir.mkdir(parents=True, exist_ok=True)
            
            logger.info(f"开始转换PDF到图片: {path}")
            
            if pages:
                page_list = [int(p.strip()) - 1 for p in pages.split(',') if p.strip().isdigit()]
                first_page = min(page_list) + 1 if page_list else 1
                last_page = max(page_list) + 1 if page_list else None
            else:
                first_page = 1
                last_page = None
                page_list = None
            
            loop = asyncio.get_event_loop()
            images = await loop.run_in_executor(
                None, 
                lambda: convert_from_path(
                    str(path),
                    dpi=dpi,
                    first_page=first_page,
                    last_page=last_page
                )
            )
            
            created_files = []
            for i, image in enumerate(images):
                page_num = first_page + i if page_list is None else page_list[i] + 1
                output_file = output_dir / f"{path.stem}_page_{page_num}.png"
                image.save(str(output_file), 'PNG')
                created_files.append(str(output_file))
            
            logger.info(f"✅ PDF转图片成功: {len(created_files)} 张")
            return f"""✅ 已转换为图片

📁 保存目录: {output_dir}
🖼️ 创建图片: {len(created_files)} 张

{chr(10).join([Path(f).name for f in created_files[:10]])}"""
        
        except Exception as e:
            logger.error(f"PDF转图片失败: {e}")
            return f"❌ 转换失败: {str(e)}"

    async def _handle_word_to_pdf(self, params: Dict) -> str:
        """将Word文档转换为PDF"""
        path = params.get("path", "") or params.get("file_path", "")
        
        if not path:
            return "❌ 请提供Word文件路径"
        
        path = Path(path)
        if not path.exists():
            return f"❌ 文件不存在: {path}"
        
        if path.suffix.lower() not in [".docx", ".doc"]:
            return f"❌ 不是Word文件: {path}"
        
        output_path = path.parent / f"{path.stem}.pdf"
        
        if output_path.exists():
            return f"⚠️ 目标文件已存在: {output_path}"
        
        logger.info(f"开始转换Word到PDF: {path} → {output_path}")
        
        try:
            import asyncio
            import win32com.client
            import pythoncom
            
            def convert_word_to_pdf():
                pythoncom.CoInitialize()
                word = None
                doc = None
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    word.DisplayAlerts = 0
                    
                    doc = word.Documents.Open(str(path.absolute()))
                    doc.SaveAs(str(output_path.absolute()), FileFormat=17)
                    doc.Close()
                    
                    logger.info(f"✅ Word转PDF成功: {output_path}")
                    return True
                except Exception as e:
                    logger.error(f"COM转换失败: {e}")
                    return False
                finally:
                    if doc:
                        try:
                            doc.Close()
                        except:
                            pass
                    if word:
                        try:
                            word.Quit()
                        except:
                            pass
                    pythoncom.CoUninitialize()
            
            loop = asyncio.get_event_loop()
            success = await loop.run_in_executor(None, convert_word_to_pdf)
            
            if success and output_path.exists():
                return f"✅ 已转换为PDF\n\n📁 保存位置: {output_path}"
            else:
                return await self._convert_word_with_libreoffice(path, output_path)
        
        except ImportError:
            return await self._convert_word_with_libreoffice(path, output_path)
        except Exception as e:
            logger.error(f"Word转PDF失败: {e}")
            return f"❌ 转换失败: {str(e)}"

    async def _convert_word_with_libreoffice(self, path: Path, output_path: Path) -> str:
        """使用 LibreOffice 转换 Word 到 PDF"""
        import asyncio
        import subprocess
        
        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        
        soffice = None
        for p in libreoffice_paths:
            if Path(p).exists():
                soffice = p
                break
        
        if not soffice:
            return """❌ Word转PDF需要以下条件之一：

1. **Microsoft Word** (Windows)
   - 确保已安装 Microsoft Office
   
2. **LibreOffice** (跨平台)
   - 下载: https://www.libreoffice.org/download/
   - 安装后重试

💡 或者使用在线转换工具"""
        
        try:
            cmd = [
                soffice,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(path.parent),
                str(path)
            ]
            
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None, 
                lambda: subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            )
            
            if output_path.exists():
                logger.info(f"✅ LibreOffice转换成功: {output_path}")
                return f"✅ 已转换为PDF\n\n📁 保存位置: {output_path}"
            else:
                return f"❌ LibreOffice转换失败: {result.stderr}"
        
        except Exception as e:
            logger.error(f"LibreOffice转换失败: {e}")
            return f"❌ 转换失败: {str(e)}"

    async def _handle_txt_to_pdf(self, params: Dict) -> str:
        """将TXT文本转换为PDF"""
        path = params.get("path", "") or params.get("file_path", "")
        
        if not path:
            return "❌ 请提供TXT文件路径"
        
        path = Path(path)
        if not path.exists():
            return f"❌ 文件不存在: {path}"
        
        if path.suffix.lower() != ".txt":
            return f"❌ 不是TXT文件: {path}"
        
        if not self._check_reportlab():
            return "❌ 需要安装 reportlab 库: pip install reportlab"
        
        try:
            content = path.read_text(encoding='utf-8')
            output_path = path.parent / f"{path.stem}.pdf"
            
            if output_path.exists():
                return f"⚠️ 目标文件已存在: {output_path}"
            
            logger.info(f"开始转换TXT到PDF: {path}")
            
            result = await self._generate_with_reportlab(path.stem, content, output_path)
            
            if "✅" in result:
                return f"✅ 已转换为PDF\n\n📁 保存位置: {output_path}"
            return result
        
        except Exception as e:
            logger.error(f"TXT转PDF失败: {e}")
            return f"❌ 转换失败: {str(e)}"

    async def _handle_image_to_pdf(self, params: Dict) -> str:
        """将图片转换为PDF"""
        path = params.get("path", "") or params.get("file_path", "")
        
        if not path:
            return "❌ 请提供图片文件路径"
        
        path = Path(path)
        if not path.exists():
            return f"❌ 文件不存在: {path}"
        
        if path.suffix.lower() not in [".png", ".jpg", ".jpeg", ".bmp", ".gif"]:
            return f"❌ 不是图片文件: {path}"
        
        try:
            from PIL import Image
            
            output_path = path.parent / f"{path.stem}.pdf"
            
            if output_path.exists():
                return f"⚠️ 目标文件已存在: {output_path}"
            
            logger.info(f"开始转换图片到PDF: {path}")
            
            img = Image.open(str(path))
            
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            
            img.save(str(output_path), 'PDF')
            
            logger.info(f"✅ 图片转PDF成功: {output_path}")
            return f"✅ 已转换为PDF\n\n📁 保存位置: {output_path}"
        
        except ImportError:
            return "❌ 需要安装 Pillow 库: pip install Pillow"
        except Exception as e:
            logger.error(f"图片转PDF失败: {e}")
            return f"❌ 转换失败: {str(e)}"

    async def _handle_excel_to_pdf(self, params: Dict) -> str:
        """将Excel文件转换为PDF"""
        path = params.get("path", "") or params.get("file_path", "")
        
        if not path:
            return "❌ 请提供Excel文件路径"
        
        path = Path(path)
        if not path.exists():
            return f"❌ 文件不存在: {path}"
        
        if path.suffix.lower() not in [".xlsx", ".xls"]:
            return f"❌ 不是Excel文件: {path}"
        
        try:
            import asyncio
            import win32com.client
            import pythoncom
            
            output_path = path.parent / f"{path.stem}.pdf"
            
            if output_path.exists():
                return f"⚠️ 目标文件已存在: {output_path}"
            
            logger.info(f"开始转换Excel到PDF: {path}")
            
            def convert_excel_to_pdf():
                pythoncom.CoInitialize()
                try:
                    excel = win32com.client.Dispatch("Excel.Application")
                    excel.Visible = False
                    
                    wb = excel.Workbooks.Open(str(path.absolute()))
                    wb.ExportAsFixedFormat(0, str(output_path.absolute()))
                    wb.Close()
                    excel.Quit()
                finally:
                    pythoncom.CoUninitialize()
            
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, convert_excel_to_pdf)
            
            logger.info(f"✅ Excel转PDF成功: {output_path}")
            return f"✅ 已转换为PDF\n\n📁 保存位置: {output_path}"
        
        except ImportError:
            return "❌ 需要安装 pywin32 库: pip install pywin32\n\n💡 注意：此方法仅支持 Windows，且需要安装 Microsoft Excel"
        except Exception as e:
            logger.error(f"Excel转PDF失败: {e}")
            return f"❌ 转换失败: {str(e)}"

    def get_status(self) -> Dict[str, Any]:
        """获取智能体状态"""
        status = super().get_status()
        status.update({
            "capabilities": list(self.capabilities),
            "pypdf_available": self._check_pypdf(),
            "reportlab_available": self._check_reportlab(),
            "python_docx_available": self._check_python_docx(),
            "openpyxl_available": self._check_openpyxl(),
            "pdf2docx_available": self._check_pdf2docx(),
            "pdf2image_available": self._check_pdf2image(),
        })
        return status
    def _get_help_info(self) -> str:
        """获取帮助信息"""
        return """## 文档智能体

### 功能说明
文档智能体可以处理各种文档操作，包括PDF转换、文档生成等。

### 支持的操作
- **PDF转Word**：将PDF转换为Word文档
- **Word转PDF**：将Word转换为PDF
- **PDF合并**：合并多个PDF文件
- **PDF拆分**：拆分PDF文件
- **文档生成**：生成各种格式文档

### 使用示例
- "把PDF转成Word" - PDF转Word
- "合并PDF文件" - 合并多个PDF
- "生成PDF文档" - 生成新文档

### 注意事项
- 支持多种文档格式
- 大文件处理可能需要时间
- 转换后的文档会保存在指定目录"""
