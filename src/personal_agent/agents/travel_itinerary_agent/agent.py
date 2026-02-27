"""
旅游攻略智能体 - 生成旅游攻略并导出为Word文档
支持根据目的地和天数生成结构化旅游攻略
"""
import re
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime
from loguru import logger

from ..base import BaseAgent, Task


class TravelItineraryAgent(BaseAgent):
    """旅游攻略智能体 - 生成旅游攻略并导出为Word文档"""
    
    KEYWORD_MAPPINGS = {
        "旅游攻略": ("generate", {}),
        "旅行攻略": ("generate", {}),
        "旅游计划": ("generate", {}),
        "旅行计划": ("generate", {}),
        "行程规划": ("generate", {}),
        "行程安排": ("generate", {}),
        "旅游路线": ("generate", {}),
        "旅行路线": ("generate", {}),
        "写攻略": ("generate", {}),
        "生成攻略": ("generate", {}),
        "制作攻略": ("generate", {}),
        "制定攻略": ("generate", {}),
    }
    
    TRAVEL_DATA = {
        "西安": {
            "description": "十三朝古都，中华文明的重要发祥地，拥有丰富的历史文化遗产",
            "attractions": {
                "必游景点": [
                    {"name": "秦始皇兵马俑博物馆", "time": "3-4小时", "price": "120元", "tips": "世界第八大奇迹，建议请导游讲解"},
                    {"name": "华清宫", "time": "2-3小时", "price": "120元", "tips": "杨贵妃沐浴之地，可观看《长恨歌》演出"},
                    {"name": "大雁塔", "time": "1-2小时", "price": "40元", "tips": "玄奘法师译经之地，晚上有音乐喷泉"},
                    {"name": "西安城墙", "time": "2-3小时", "price": "54元", "tips": "中国现存最完整的古城墙，可骑行"},
                    {"name": "陕西历史博物馆", "time": "2-3小时", "price": "免费", "tips": "需提前预约，馆藏珍宝无数"},
                    {"name": "回民街", "time": "2小时", "price": "免费", "tips": "美食天堂，羊肉泡馍必尝"},
                    {"name": "钟鼓楼", "time": "1小时", "price": "30元", "tips": "西安地标，夜景更美"},
                    {"name": "大唐不夜城", "time": "2-3小时", "price": "免费", "tips": "网红打卡地，晚上灯光璀璨"},
                ],
                "深度景点": [
                    {"name": "小雁塔", "time": "1-2小时", "price": "免费", "tips": "人少清静，适合拍照"},
                    {"name": "碑林博物馆", "time": "2小时", "price": "50元", "tips": "书法爱好者必去"},
                    {"name": "大明宫遗址", "time": "2-3小时", "price": "60元", "tips": "唐朝皇宫遗址，气势恢宏"},
                ]
            },
            "foods": [
                {"name": "羊肉泡馍", "recommend": "老孙家、同盛祥", "price": "30-50元"},
                {"name": "肉夹馍", "recommend": "樊记肉夹馍、秦豫肉夹馍", "price": "15-25元"},
                {"name": "凉皮", "recommend": "魏家凉皮、盛志望麻酱酿皮", "price": "10-20元"},
                {"name": "biangbiang面", "recommend": "老碗、长安大排档", "price": "20-35元"},
                {"name": "胡辣汤", "recommend": "老金家、刘老虎", "price": "10-15元"},
                {"name": "甑糕", "recommend": "东南亚甑糕", "price": "10-15元"},
                {"name": "灌汤包", "recommend": "贾三灌汤包", "price": "25-40元"},
            ],
            "hotels": {
                "钟楼附近": "交通便利，靠近回民街，适合首次来西安",
                "大雁塔附近": "环境优美，靠近大唐不夜城，适合休闲游",
                "城墙附近": "古色古香，体验古城风情",
            },
            "transport": "地铁2号线贯穿南北，可达大部分景点；机场大巴直达市区；出租车起步价10元",
            "tips": [
                "兵马俑建议早上8点前到达，避开人流高峰",
                "陕西历史博物馆需提前在官网预约",
                "回民街美食众多，注意选择正规店铺",
                "城墙骑行建议傍晚时分，可欣赏夕阳",
                "大唐不夜城晚上7点后灯光最美",
            ]
        },
        "北京": {
            "description": "中国首都，拥有3000多年建城史，世界文化遗产众多",
            "attractions": {
                "必游景点": [
                    {"name": "故宫博物院", "time": "4-5小时", "price": "60元", "tips": "需提前10天预约，周一闭馆"},
                    {"name": "长城（八达岭）", "time": "半天", "price": "40元", "tips": "建议乘坐S2线，避开节假日"},
                    {"name": "天安门广场", "time": "1-2小时", "price": "免费", "tips": "升旗仪式值得观看"},
                    {"name": "颐和园", "time": "3-4小时", "price": "30元", "tips": "皇家园林，四季皆美"},
                    {"name": "天坛公园", "time": "2-3小时", "price": "15元", "tips": "明清皇帝祭天之地"},
                    {"name": "南锣鼓巷", "time": "2小时", "price": "免费", "tips": "老北京胡同文化"},
                    {"name": "鸟巢/水立方", "time": "1-2小时", "price": "50元", "tips": "奥运场馆，夜景更美"},
                ]
            },
            "foods": [
                {"name": "北京烤鸭", "recommend": "全聚德、便宜坊", "price": "100-200元"},
                {"name": "炸酱面", "recommend": "老北京炸酱面", "price": "20-30元"},
                {"name": "豆汁焦圈", "recommend": "护国寺小吃", "price": "10-15元"},
                {"name": "涮羊肉", "recommend": "东来顺", "price": "80-150元"},
            ],
            "hotels": {
                "王府井": "购物方便，靠近故宫",
                "前门": "老北京风情，交通便利",
                "三里屯": "时尚商圈，夜生活丰富",
            },
            "transport": "地铁网络发达，覆盖主要景点；公交IC卡通用",
            "tips": [
                "故宫门票需提前在官网预约",
                "长城建议工作日前往",
                "北京冬天较冷，注意保暖",
            ]
        },
        "上海": {
            "description": "国际化大都市，东西方文化交融的现代都市",
            "attractions": {
                "必游景点": [
                    {"name": "外滩", "time": "2小时", "price": "免费", "tips": "万国建筑博览群，夜景绝美"},
                    {"name": "东方明珠", "time": "2-3小时", "price": "180元", "tips": "上海地标，可俯瞰全城"},
                    {"name": "豫园", "time": "2小时", "price": "40元", "tips": "江南古典园林"},
                    {"name": "南京路步行街", "time": "2小时", "price": "免费", "tips": "中华商业第一街"},
                    {"name": "田子坊", "time": "2小时", "price": "免费", "tips": "文艺小资聚集地"},
                    {"name": "迪士尼乐园", "time": "全天", "price": "399元起", "tips": "建议提前购票，工作日人少"},
                ]
            },
            "foods": [
                {"name": "小笼包", "recommend": "南翔馒头店、鼎泰丰", "price": "30-50元"},
                {"name": "生煎包", "recommend": "小杨生煎", "price": "15-25元"},
                {"name": "白切鸡", "recommend": "小绍兴", "price": "50-80元"},
                {"name": "红烧肉", "recommend": "老正兴", "price": "60-100元"},
            ],
            "hotels": {
                "外滩": "景观绝佳，价格较高",
                "南京路": "购物方便，交通便利",
                "陆家嘴": "现代商务区，高端酒店多",
            },
            "transport": "地铁覆盖全城，磁悬浮列车连接机场",
            "tips": [
                "外滩夜景最佳观赏时间为晚上7-9点",
                "迪士尼建议工作日前往",
                "上海夏季炎热，注意防暑",
            ]
        },
        "成都": {
            "description": "天府之国，美食之都，大熊猫的故乡",
            "attractions": {
                "必游景点": [
                    {"name": "大熊猫繁育研究基地", "time": "3-4小时", "price": "55元", "tips": "建议早上8点前到达"},
                    {"name": "宽窄巷子", "time": "2-3小时", "price": "免费", "tips": "成都慢生活代表"},
                    {"name": "锦里古街", "time": "2小时", "price": "免费", "tips": "三国文化与美食结合"},
                    {"name": "武侯祠", "time": "2小时", "price": "50元", "tips": "诸葛亮纪念地"},
                    {"name": "杜甫草堂", "time": "2小时", "price": "50元", "tips": "诗圣故居"},
                    {"name": "春熙路", "time": "2小时", "price": "免费", "tips": "成都最繁华商业街"},
                ]
            },
            "foods": [
                {"name": "火锅", "recommend": "小龙坎、蜀大侠", "price": "80-150元"},
                {"name": "串串香", "recommend": "钢管厂五区", "price": "40-60元"},
                {"name": "担担面", "recommend": "小名堂", "price": "15-25元"},
                {"name": "龙抄手", "recommend": "龙抄手总店", "price": "20-30元"},
                {"name": "兔头", "recommend": "双流老妈兔头", "price": "15-25元"},
            ],
            "hotels": {
                "春熙路": "购物方便，交通便利",
                "宽窄巷子": "体验成都慢生活",
                "锦里附近": "靠近景点，古色古香",
            },
            "transport": "地铁覆盖主要景点，打车方便",
            "tips": [
                "看熊猫要早起，熊猫上午最活跃",
                "成都美食偏辣，注意肠胃",
                "宽窄巷子适合下午茶时光",
            ]
        },
        "杭州": {
            "description": "人间天堂，西湖美景闻名天下",
            "attractions": {
                "必游景点": [
                    {"name": "西湖", "time": "半天", "price": "免费", "tips": "建议骑行或步行环湖"},
                    {"name": "灵隐寺", "time": "2-3小时", "price": "75元", "tips": "江南名刹，香火鼎盛"},
                    {"name": "雷峰塔", "time": "1-2小时", "price": "40元", "tips": "白娘子传说发源地"},
                    {"name": "西溪湿地", "time": "3-4小时", "price": "80元", "tips": "城市绿肺，适合休闲"},
                    {"name": "河坊街", "time": "2小时", "price": "免费", "tips": "古街美食众多"},
                    {"name": "宋城", "time": "半天", "price": "310元", "tips": "《宋城千古情》必看"},
                ]
            },
            "foods": [
                {"name": "西湖醋鱼", "recommend": "楼外楼", "price": "80-120元"},
                {"name": "东坡肉", "recommend": "知味观", "price": "30-50元"},
                {"name": "龙井虾仁", "recommend": "楼外楼", "price": "60-100元"},
                {"name": "片儿川", "recommend": "菊英面馆", "price": "20-30元"},
            ],
            "hotels": {
                "西湖边": "景观绝佳，价格较高",
                "武林广场": "市中心，交通便利",
                "河坊街": "古色古香，美食众多",
            },
            "transport": "地铁覆盖主要景点，公共自行车方便",
            "tips": [
                "西湖最佳游览季节为春秋",
                "灵隐寺建议早上前往",
                "宋城演出需提前购票",
            ]
        },
        "大同": {
            "description": "中国九大古都之一，北魏首都，云冈石窟所在地，历史文化名城",
            "attractions": {
                "必游景点": [
                    {"name": "云冈石窟", "time": "3-4小时", "price": "120元", "tips": "世界文化遗产，建议请导游讲解，第5窟和第6窟最精彩"},
                    {"name": "悬空寺", "time": "2-3小时", "price": "130元", "tips": "建在悬崖上的古寺，恐高者慎入，建议早到避开人流"},
                    {"name": "华严寺", "time": "2小时", "price": "65元", "tips": "辽金建筑代表，薄伽教藏殿内的合掌露齿菩萨最美"},
                    {"name": "九龙壁", "time": "1小时", "price": "10元", "tips": "中国三大九龙壁之一，比北京故宫的还早"},
                    {"name": "大同古城墙", "time": "2小时", "price": "免费", "tips": "可骑行或步行，夜景灯光很美"},
                    {"name": "善化寺", "time": "1-2小时", "price": "50元", "tips": "辽金建筑完整保存，三圣殿彩塑精美"},
                ],
                "深度景点": [
                    {"name": "恒山", "time": "半天", "price": "45元", "tips": "五岳之一，可登顶远眺"},
                    {"name": "土林", "time": "2小时", "price": "50元", "tips": "地质奇观，适合摄影"},
                    {"name": "火山群", "time": "半天", "price": "免费", "tips": "大同火山群国家地质公园"},
                ]
            },
            "foods": [
                {"name": "刀削面", "recommend": "东方削面、老柴削面", "price": "15-25元"},
                {"name": "羊杂", "recommend": "帅府街羊杂", "price": "20-30元"},
                {"name": "黄糕", "recommend": "大同老字号", "price": "10-20元"},
                {"name": "浑源凉粉", "recommend": "浑源县城", "price": "10-15元"},
                {"name": "烧麦", "recommend": "凤临阁", "price": "30-50元"},
                {"name": "兔头", "recommend": "帅府街", "price": "15-25元"},
            ],
            "hotels": {
                "古城内": "靠近景点，体验古城风情",
                "火车站附近": "交通便利，价格实惠",
                "云冈附近": "方便游览石窟，环境清静",
            },
            "transport": "市区公交便利，去云冈石窟可乘旅游专线；去悬空寺建议包车或跟团",
            "tips": [
                "云冈石窟建议上午前往，光线较好",
                "悬空寺距离市区较远，需预留半天时间",
                "大同早晚温差大，注意保暖",
                "品尝刀削面要趁热，口感最佳",
            ]
        }
    }
    
    def __init__(self):
        super().__init__(
            name="travel_itinerary_agent",
            description="旅游攻略智能体 - 生成旅游攻略并导出为Word文档"
        )
        
        self.register_capability(
            capability="create_travel_plan",
            description="生成旅游攻略。根据目的地和天数生成详细的旅游行程规划。",
            parameters={
                "type": "object",
                "properties": {
                    "destination": {
                        "type": "string",
                        "description": "旅游目的地"
                    },
                    "days": {
                        "type": "integer",
                        "description": "旅游天数",
                        "default": 3
                    }
                },
                "required": ["destination"]
            },
            category="travel"
        )
        
        self.register_capability("generate_travel_itinerary", "生成旅游攻略")
        self.register_capability("export_to_docx", "导出为Word文档")
        
        self._llm_gateway = None
        
        logger.info("🗺️ 旅游攻略智能体已初始化")

    def _get_llm_gateway(self):
        """获取 LLM 网关"""
        if self._llm_gateway is None:
            from ...llm import LLMGateway
            from ...config import settings
            self._llm_gateway = LLMGateway(settings.llm)
        return self._llm_gateway

    async def _call_llm(self, prompt: str) -> str:
        """调用 LLM"""
        llm = self._get_llm_gateway()
        messages = [{"role": "user", "content": prompt}]
        response = await llm.chat(messages)
        return response.content

    async def execute_task(self, task: Task) -> Any:
        task_type = task.type
        params = task.params
        logger.info(f"🗺️ 旅游攻略智能体执行任务: {task_type}")
        
        if task_type == "action":
            return await self._handle_action(params)
        elif task_type in ["generate", "generate_travel_itinerary", "create_travel_plan"]:
            return await self._handle_generate(params)
        elif task_type == "export_to_docx":
            return await self._handle_export(params)
        elif task_type == "agent_help":
            return self._get_help_info()
        else:
            return self.cannot_handle(f"不支持的任务类型: {task_type}")

    async def _handle_action(self, params: Dict) -> str:
        """处理通用操作"""
        text = params.get("text", "") or params.get("original_text", "")
        return await self._generate_itinerary(text)

    async def _handle_generate(self, params: Dict) -> str:
        """生成旅游攻略"""
        destination = params.get("destination", "") or params.get("city", "")
        days = params.get("days", 3)
        original_text = params.get("original_text", "")
        
        if not destination and original_text:
            destination, days = self._parse_destination_and_days(original_text)
        
        if not destination:
            return "❌ 请提供目的地城市名称，如'西安三天旅游攻略'"
        
        return await self._generate_itinerary(destination, days)

    async def _handle_export(self, params: Dict) -> str:
        """导出为Word文档"""
        content = params.get("content", "")
        title = params.get("title", "旅游攻略")
        
        if not content:
            return "❌ 请提供攻略内容"
        
        return await self._export_to_docx(content, title)

    def _parse_destination_and_days(self, text: str) -> tuple:
        """从文本中解析目的地和天数"""
        destination = ""
        days = 3
        
        day_patterns = [
            (r'(\d+)天', lambda m: int(m.group(1))),
            (r'(\d+)日', lambda m: int(m.group(1))),
            (r'三天', lambda m: 3),
            (r'两天', lambda m: 2),
            (r'一天', lambda m: 1),
            (r'四天', lambda m: 4),
            (r'五天', lambda m: 5),
        ]
        
        for pattern, extractor in day_patterns:
            match = re.search(pattern, text)
            if match:
                days = extractor(match)
                break
        
        city_patterns = [
            r'([北京上海广州深圳杭州南京成都重庆武汉西安苏州郑州长沙沈阳青岛合肥佛山宁波昆明福州厦门哈尔滨济南大连长春太原贵阳南宁南昌石家庄兰州银川西宁海口三亚拉萨呼和浩特乌鲁木齐大同]+)',
            r'去(.+?)(?:旅游|旅行|玩)',
            r'(.+?)(?:旅游|旅行)攻略',
            r'(.+?)(?:三天|两天|一天|四天|五天|\d+天|\d+日)',
        ]
        
        for pattern in city_patterns:
            match = re.search(pattern, text)
            if match:
                candidate = match.group(1).strip()
                if candidate in self.TRAVEL_DATA:
                    destination = candidate
                    break
                for city in self.TRAVEL_DATA.keys():
                    if city in candidate or candidate in city:
                        destination = city
                        break
                if destination:
                    break
        
        if not destination:
            for city in self.TRAVEL_DATA.keys():
                if city in text:
                    destination = city
                    break
        
        return destination, days

    async def _generate_itinerary(self, destination: str, days: int = 3) -> str:
        """生成旅游攻略"""
        city_data = self.TRAVEL_DATA.get(destination)
        
        if not city_data:
            content = await self._generate_with_llm(destination, days)
        else:
            content = self._build_itinerary(destination, city_data, days)
        
        return content

    def _build_itinerary(self, destination: str, city_data: Dict, days: int) -> str:
        """构建旅游攻略内容"""
        lines = []
        
        lines.append(f"# {destination}{days}天旅游攻略")
        lines.append("")
        lines.append(f"**{city_data['description']}**")
        lines.append("")
        
        attractions = city_data.get("attractions", {})
        must_visit = attractions.get("必游景点", [])
        deep_visit = attractions.get("深度景点", [])
        
        all_attractions = must_visit + deep_visit
        
        daily_attractions = []
        attractions_per_day = max(2, len(all_attractions) // days)
        
        for i in range(days):
            start_idx = i * attractions_per_day
            end_idx = start_idx + attractions_per_day
            daily_attractions.append(all_attractions[start_idx:end_idx])
        
        if len(all_attractions) > days * attractions_per_day:
            remaining = all_attractions[days * attractions_per_day:]
            for i, attr in enumerate(remaining):
                if i < len(daily_attractions):
                    daily_attractions[i].append(attr)
        
        for day_num in range(days):
            lines.append(f"## 第{day_num + 1}天")
            lines.append("")
            
            day_attractions = daily_attractions[day_num] if day_num < len(daily_attractions) else []
            
            if day_attractions:
                for i, attr in enumerate(day_attractions):
                    lines.append(f"### {i + 1}. {attr['name']}")
                    lines.append("")
                    lines.append(f"- **游览时间**: {attr['time']}")
                    lines.append(f"- **门票价格**: {attr['price']}")
                    lines.append(f"- **游玩建议**: {attr['tips']}")
                    lines.append("")
            else:
                lines.append("自由活动，可根据个人兴趣安排行程。")
                lines.append("")
        
        lines.append("## 美食推荐")
        lines.append("")
        foods = city_data.get("foods", [])
        for food in foods:
            lines.append(f"- **{food['name']}**: 推荐店铺 - {food['recommend']}，人均 {food['price']}")
        lines.append("")
        
        lines.append("## 住宿建议")
        lines.append("")
        hotels = city_data.get("hotels", {})
        for area, desc in hotels.items():
            lines.append(f"- **{area}**: {desc}")
        lines.append("")
        
        lines.append("## 交通指南")
        lines.append("")
        lines.append(city_data.get("transport", "建议使用公共交通出行"))
        lines.append("")
        
        lines.append("## 实用贴士")
        lines.append("")
        tips = city_data.get("tips", [])
        for tip in tips:
            lines.append(f"- {tip}")
        lines.append("")
        
        lines.append("---")
        lines.append(f"*攻略生成时间: {datetime.now().strftime('%Y年%m月%d日')}*")
        
        return "\n".join(lines)

    async def _generate_with_llm(self, destination: str, days: int) -> str:
        """使用LLM生成旅游攻略"""
        prompt = f"""请为{destination}生成一份{days}天的旅游攻略，包含以下内容：

1. 城市简介
2. 每日行程安排（包含景点名称、游览时间、门票价格、游玩建议）
3. 美食推荐（特色美食、推荐店铺、人均消费）
4. 住宿建议（推荐区域及特点）
5. 交通指南
6. 实用贴士

请用Markdown格式输出，结构清晰，内容实用。"""

        try:
            content = await self._call_llm(prompt)
            return content
        except Exception as e:
            logger.error(f"LLM生成攻略失败: {e}")
            return f"# {destination}{days}天旅游攻略\n\n抱歉，暂时无法生成详细攻略，请稍后再试。"

    async def _export_to_docx(self, content: str, title: str) -> str:
        """导出为Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re
            
            from ...config import settings
            output_dir = settings.directory.get_documents_dir() / "旅游攻略"
            output_dir.mkdir(parents=True, exist_ok=True)
            
            safe_title = re.sub(r'[\\/:*?"<>|]', '', title)
            output_path = output_dir / f"{safe_title}.docx"
            
            doc = Document()
            
            style = doc.styles['Normal']
            font = style.font
            font.name = 'SimSun'
            font.size = Pt(12)
            
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            lines = content.split('\n')
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    continue
                
                if stripped.startswith('# ') and not stripped.startswith('## '):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith('## ') and not stripped.startswith('### '):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith('### '):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith('- ') or stripped.startswith('* '):
                    text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped[2:])
                    doc.add_paragraph(text, style='List Bullet')
                elif re.match(r'^\d+\.\s', stripped):
                    text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
                    doc.add_paragraph(text, style='List Number')
                elif stripped.startswith('**') and stripped.endswith('**'):
                    para = doc.add_paragraph()
                    run = para.add_run(stripped[2:-2])
                    run.bold = True
                else:
                    text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
                    doc.add_paragraph(text)
            
            doc.save(str(output_path))
            
            logger.info(f"✅ 旅游攻略Word文档已生成: {output_path}")
            return f"📁 文件已保存: {output_path}\n\n💡 可以用 Word 或 WPS 打开查看"
            
        except ImportError:
            logger.warning("python-docx 未安装，尝试使用 document_agent")
            return await self._export_via_document_agent(content, title)
        except Exception as e:
            logger.error(f"导出Word文档失败: {e}")
            return f"❌ 导出失败: {e}"

    async def _export_via_document_agent(self, content: str, title: str) -> str:
        """通过 document_agent 导出"""
        try:
            from ..multi_agent_system import multi_agent_system
            
            if multi_agent_system and multi_agent_system.master:
                doc_agent = multi_agent_system.master.sub_agents.get('document_agent')
                if doc_agent:
                    task = Task(
                        type="doc_generate",
                        content=title,
                        params={"content": content, "title": title}
                    )
                    result = await doc_agent.execute_task(task)
                    return result
        except Exception as e:
            logger.error(f"通过 document_agent 导出失败: {e}")
        
        return "❌ 导出失败，请确保已安装 python-docx 库: pip install python-docx"
